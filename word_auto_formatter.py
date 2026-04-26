import os
import re
import shutil
import uuid
from typing import Dict, Optional, Set

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from word_format_checker import CN_HEITI, CN_SIMSUN, DEFAULT_STANDARDS, _analyze_document_structure, _merge, detect_word_format


MIN_FIX_CONFIDENCE = 0.68


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False):
    run.font.name = font_name
    if run._element is not None and run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def _parse_paragraph_index(position: str) -> Optional[int]:
    value = str(position or "").strip()
    if not value:
        return None
    value = value.translate(str.maketrans("０１２３４５６７８９", "0123456789"))
    match = re.search(r"paragraph\s+(\d+)", value, flags=re.IGNORECASE)
    if match:
        return int(match.group(1))
    match = re.search(r"第\s*(\d+)\s*段", value)
    if match:
        return int(match.group(1))
    match = re.search(r"段落\s*(\d+)", value)
    return int(match.group(1)) if match else None


def _issue_status(issue: Dict) -> str:
    status = str(issue.get("issue_status", "open") or "open").strip().lower()
    if status == "open":
        confidence = float(issue.get("confidence", 0.0) or 0.0)
        reasons = [str(item).lower() for item in issue.get("reasons", [])]
        if confidence >= MIN_FIX_CONFIDENCE and "suspected_heading" not in reasons:
            return "confirmed"
    return status


def _is_fixable_issue(issue: Dict) -> bool:
    confidence = float(issue.get("confidence", 0.0) or 0.0)
    if confidence < MIN_FIX_CONFIDENCE:
        return False

    problem_type = str(issue.get("problem_type", "") or "")
    region = str(issue.get("region", "") or "")
    para_type = str(issue.get("paragraph_type", "") or "")
    status = _issue_status(issue)
    safe_manual_review_fixables = {
        "body_font_unknown",
        "body_font_size_unknown",
        "body_line_spacing_unknown",
        "body_first_indent_unknown",
        "heading_font_unknown",
        "heading_font_size_unknown",
        "heading_bold_unknown",
        "heading_alignment_unknown",
    }
    if status != "confirmed":
        if not (
            status == "manual_review"
            and problem_type in safe_manual_review_fixables
            and (
                (region == "body" and para_type == "body")
                or (region in {"body", "appendix"} and para_type in {"heading_1", "heading_2", "heading_3"})
            )
        ):
            return False

    reasons = [str(item).lower() for item in issue.get("reasons", [])]
    if any(tag in reasons for tag in ("suspected_heading", "unsupported_reference_number_format")):
        return False

    non_fixable = {
        "suspected_heading",
        "heading_number",
        "heading_hierarchy",
        "reference_numbering",
        "reference_number_duplicate",
        "reference_number_not_continuous",
        "reference_number_out_of_order",
        "reference_numbering_mixed_format",
        "reference_entry_empty",
        "reference_entry_unrecognized",
        "reference_entries_uncertain",
        "reference_region_missing",
        "reference_region_uncertain",
        "toc",
        "page_number",
    }
    return problem_type not in non_fixable


def _set_page_layout(doc: Document):
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def _style_body(paragraph, body_standard: Dict):
    font_name = str(body_standard.get("font_name", CN_SIMSUN))
    font_size_pt = float(body_standard.get("font_size_pt", 12.0))
    line_spacing = float(body_standard.get("line_spacing", 1.5))
    first_line_indent_chars = float(body_standard.get("first_line_indent_chars", 2.0))
    space_before_pt = float(body_standard.get("space_before_pt", 0.0))
    space_after_pt = float(body_standard.get("space_after_pt", 0.0))

    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.first_line_indent = Pt(font_size_pt * first_line_indent_chars)
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        if run.text and run.text.strip():
            _set_run_font(run, font_name, font_size_pt, bold=False)


def _style_heading(paragraph, level: int, headings_standard: Dict):
    heading = headings_standard.get(f"h{level}", {}) if isinstance(headings_standard, dict) else {}
    align_name = str(heading.get("align", "left")).lower()
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    font_name = str(heading.get("font_name", CN_HEITI))
    size_pt = float(heading.get("font_size_pt", 16.0 if level == 1 else 14.0 if level == 2 else 12.0))
    bold = bool(heading.get("bold", True))
    align = align_map.get(align_name, WD_ALIGN_PARAGRAPH.LEFT)

    if level == 1:
        align = align_map.get(align_name, WD_ALIGN_PARAGRAPH.CENTER)

    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.alignment = align
    for run in paragraph.runs:
        if run.text and run.text.strip():
            _set_run_font(run, font_name, size_pt, bold=bold)


def _style_caption(paragraph):
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    for run in paragraph.runs:
        if run.text and run.text.strip():
            _set_run_font(run, CN_SIMSUN, 10.5, bold=False)


def _style_figure_caption(paragraph):
    _style_caption(paragraph)


def _style_table_caption(paragraph):
    _style_caption(paragraph)


def _style_reference_item(paragraph):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.left_indent = Pt(24)
    paragraph.paragraph_format.first_line_indent = Pt(-24)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        if run.text and run.text.strip():
            _set_run_font(run, CN_SIMSUN, 10.5, bold=False)


def _style_image_paragraph(paragraph):
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _safe_for_action(ctx: Dict, action: str) -> bool:
    if not ctx:
        return False
    if float(ctx.get("confidence", 0.0) or 0.0) < MIN_FIX_CONFIDENCE:
        return False
    if ctx.get("issue_status") in {"suspected", "manual_review"}:
        return False
    if ctx.get("paragraph_type") == "unknown":
        return False

    region = ctx.get("region")
    para_type = ctx.get("paragraph_type")

    if action == "body":
        return region == "body" and para_type == "body"
    if action == "heading":
        return region in {"body", "appendix"} and para_type in {"heading_1", "heading_2", "heading_3"}
    if action == "figure_caption":
        return para_type == "figure_caption"
    if action == "table_caption":
        return para_type == "table_caption"
    if action == "reference_item":
        return region == "references" and para_type == "reference_item"
    if action == "image_paragraph":
        return para_type != "empty"
    return False


def _issue_to_action(issue: Dict) -> Optional[str]:
    problem_type = str(issue.get("problem_type", "") or "")
    para_type = str(issue.get("paragraph_type", "") or "")

    if problem_type in {
        "body_font",
        "body_font_size",
        "line_spacing",
        "first_line_indent",
        "paragraph_spacing",
        "body_font_unknown",
        "body_font_size_unknown",
        "body_line_spacing_unknown",
        "body_first_indent_unknown",
    }:
        return "body"
    if problem_type in {"heading_format", "heading_font_unknown", "heading_font_size_unknown", "heading_bold_unknown", "heading_alignment_unknown"}:
        return "heading"
    if problem_type in {"figure_caption_style"} or para_type == "figure_caption":
        return "figure_caption"
    if problem_type in {"image_alignment"}:
        return "image_paragraph"
    if problem_type in {"table_caption_style"} or para_type == "table_caption":
        return "table_caption"
    if problem_type in {"reference_item_style"} or para_type == "reference_item":
        return "reference_item"
    return None


def _safe_output_path(output_dir: str, prefix: str) -> Dict[str, str]:
    os.makedirs(output_dir, exist_ok=True)
    while True:
        output_name = f"{prefix}_{uuid.uuid4().hex[:12]}.docx"
        output_path = os.path.join(output_dir, output_name)
        if not os.path.exists(output_path):
            return {"output_name": output_name, "output_path": output_path}


def auto_format_docx(
    input_path: str,
    output_dir: str,
    prefix: str = "formatted",
    standards: Optional[Dict] = None,
    checks: Optional[Dict[str, bool]] = None,
    severities: Optional[Dict[str, str]] = None,
) -> Dict:
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"source file not found: {input_path}")

    out_meta = _safe_output_path(output_dir, prefix)
    output_name = out_meta["output_name"]
    output_path = out_meta["output_path"]

    # Safety strategy: copy source first and only edit the copied file.
    shutil.copy2(input_path, output_path)

    final_standards = _merge(DEFAULT_STANDARDS, standards or {})
    final_checks = checks or {}
    final_severities = severities or {}

    check_result = detect_word_format(output_path, standards=standards or {}, checks=final_checks, severities=final_severities)
    issues = list(check_result.get("issues", []))

    doc = Document(output_path)
    contexts = _analyze_document_structure(doc, final_standards)
    context_by_idx = {ctx["index"]: ctx for ctx in contexts}

    need_page_layout = False
    paragraph_actions: Dict[int, Set[str]] = {}
    skipped_issue_count = 0
    confirmed_issue_count = 0

    for issue in issues:
        if not _is_fixable_issue(issue):
            skipped_issue_count += 1
            continue
        confirmed_issue_count += 1

        problem_type = str(issue.get("problem_type", "") or "")
        if problem_type in {"page_margin", "paper_size"}:
            need_page_layout = True
            continue

        paragraph_idx = _parse_paragraph_index(issue.get("position", ""))
        if not paragraph_idx:
            skipped_issue_count += 1
            continue
        if paragraph_idx < 1 or paragraph_idx > len(doc.paragraphs):
            skipped_issue_count += 1
            continue

        action = _issue_to_action(issue)
        if action is None:
            skipped_issue_count += 1
            continue

        ctx = context_by_idx.get(paragraph_idx)
        if not _safe_for_action(ctx, action):
            skipped_issue_count += 1
            continue

        paragraph_actions.setdefault(paragraph_idx, set()).add(action)

    if need_page_layout:
        _set_page_layout(doc)

    for idx in sorted(paragraph_actions.keys()):
        paragraph = doc.paragraphs[idx - 1]
        ctx = context_by_idx.get(idx) or {}
        actions = paragraph_actions[idx]

        if "body" in actions and _safe_for_action(ctx, "body"):
            _style_body(paragraph, final_standards.get("body", {}))

        if "heading" in actions and _safe_for_action(ctx, "heading"):
            level = int(str(ctx["paragraph_type"]).split("_")[-1])
            _style_heading(paragraph, level, final_standards.get("headings", {}))

        if "figure_caption" in actions and _safe_for_action(ctx, "figure_caption"):
            _style_figure_caption(paragraph)

        if "table_caption" in actions and _safe_for_action(ctx, "table_caption"):
            _style_table_caption(paragraph)

        if "reference_item" in actions and _safe_for_action(ctx, "reference_item"):
            _style_reference_item(paragraph)

        if "image_paragraph" in actions and _safe_for_action(ctx, "image_paragraph"):
            if "graphicData" in paragraph._p.xml:
                _style_image_paragraph(paragraph)

    doc.save(output_path)
    after_check_result = detect_word_format(output_path, standards=standards or {}, checks=final_checks, severities=final_severities)
    before_issue_count = len(issues)
    after_issue_count = len(after_check_result.get("issues", []))

    return {
        "output_name": output_name,
        "output_path": output_path,
        "file_size": os.path.getsize(output_path),
        "copied_from": input_path,
        "confirmed_issue_count": confirmed_issue_count,
        "skipped_issue_count": skipped_issue_count,
        "before_issue_count": before_issue_count,
        "after_issue_count": after_issue_count,
        "reduced_issue_count": before_issue_count - after_issue_count,
    }
