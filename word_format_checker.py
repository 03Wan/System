import json
import re
from typing import Dict, List, Optional, Set, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


CN_SIMSUN = "\u5b8b\u4f53"
CN_HEITI = "\u9ed1\u4f53"
CN_TITLE = "\u6807\u9898"

KW_ABSTRACT = "\u6458\u8981"
KW_CN_ABSTRACT = "\u4e2d\u6587\u6458\u8981"
KW_EN_ABSTRACT = "\u82f1\u6587\u6458\u8981"
KW_CATALOG = "\u76ee\u5f55"
KW_REFERENCES = "\u53c2\u8003\u6587\u732e"
KW_THANKS = "\u81f4\u8c22"
KW_ACK = "\u8c22\u8f9e"
KW_APPENDIX = "\u9644\u5f55"
KW_CHAPTER = "\u7ae0"
KW_FIGURE = "\u56fe"
KW_TABLE = "\u8868"


DEFAULT_STANDARDS = {
    "page_margin_cm": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17},
    "page": {
        "paper_width_cm": 21.0,
        "paper_height_cm": 29.7,
        "require_header": False,
        "require_footer": True,
        "page_number_position": "bottom_center",
        "page_number_format": "arabic",
        "max_sections": 1,
    },
    "body": {
        "font_name": CN_SIMSUN,
        "font_size_pt": 12.0,
        "line_spacing": 1.5,
        "first_line_indent_chars": 2,
        "align": "justify",
        "space_before_pt": 0.0,
        "space_after_pt": 0.0,
    },
    "headings": {
        "h1": {"font_name": CN_HEITI, "font_size_pt": 16.0, "bold": True, "align": "center"},
        "h2": {"font_name": CN_HEITI, "font_size_pt": 14.0, "bold": True, "align": "left"},
        "h3": {"font_name": CN_HEITI, "font_size_pt": 12.0, "bold": True, "align": "left"},
    },
    "figure_table": {
        "check_caption_position": True,
        "check_number_continuity": True,
        "check_image_center": True,
    },
    "toc": {"require_auto_toc": True},
    "references": {"continuous": True, "check_completeness": True, "check_type_mark": True, "check_punctuation": True},
    "section_titles": {
        "abstract_zh": {"font_name": CN_HEITI, "font_size_pt": 16.0, "bold": True, "align": "center"},
        "abstract_en": {"font_name": "Times New Roman", "font_size_pt": 16.0, "bold": True, "align": "center"},
        "conclusion": {"font_name": CN_HEITI, "font_size_pt": 16.0, "bold": True, "align": "center"},
        "thanks": {"font_name": CN_HEITI, "font_size_pt": 16.0, "bold": True, "align": "center"},
        "references": {"font_name": CN_HEITI, "font_size_pt": 16.0, "bold": True, "align": "left"},
        "appendix": {"font_name": CN_HEITI, "font_size_pt": 16.0, "bold": True, "align": "left"},
    },
}


DEFAULT_CHECKS = {
    "PAGE_MARGIN": True,
    "PAGE_SIZE": True,
    "HEADER_FOOTER": True,
    "PAGE_NUMBER": True,
    "SECTION_BREAK": True,
    "HEADING_LEVEL_1": True,
    "HEADING_LEVEL_2": True,
    "HEADING_LEVEL_3": True,
    "HEADING_NUMBER": True,
    "BODY_FONT": True,
    "BODY_FONT_SIZE": True,
    "BODY_LINE_SPACING": True,
    "BODY_FIRST_INDENT": True,
    "BODY_ALIGNMENT": True,
    "BODY_PARAGRAPH_SPACING": True,
    "FIGURE_TABLE": True,
    "TOC_CHECK": True,
    "REFERENCE_SEQ": True,
    "REFERENCE_FORMAT": True,
    "ABSTRACT_ZH_TITLE": True,
    "ABSTRACT_EN_TITLE": True,
    "CONCLUSION_TITLE": True,
    "THANKS_TITLE": True,
    "REFERENCE_TITLE": True,
    "APPENDIX_TITLE": True,
}


REGIONS = ("cover", "abstract", "catalog", "body", "references", "thanks", "appendix", "unknown")
PARAGRAPH_TYPES = (
    "empty",
    "body",
    "heading_1",
    "heading_2",
    "heading_3",
    "figure_caption",
    "table_caption",
    "reference_item",
    "section_marker",
    "unknown",
)

HEADING_HIGH_CONFIDENCE = 0.76
HEADING_LOW_CONFIDENCE = 0.56

CN_NUM = "\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341"
CN_CHAPTER_KEYWORDS = (
    "\u7eea\u8bba",
    "\u5f15\u8a00",
    "\u80cc\u666f",
    "\u76f8\u5173\u5de5\u4f5c",
    "\u65b9\u6cd5",
    "\u7b97\u6cd5",
    "\u5b9e\u9a8c",
    "\u7ed3\u679c",
    "\u5206\u6790",
    "\u7ed3\u8bba",
    "\u603b\u7ed3",
    "\u7cfb\u7edf\u8bbe\u8ba1",
    "\u7cfb\u7edf\u5b9e\u73b0",
)


def _pt(value) -> Optional[float]:
    if value is None:
        return None
    try:
        return float(value.pt)
    except Exception:
        return None


def _cm(value) -> Optional[float]:
    if value is None:
        return None
    try:
        return float(value.cm)
    except Exception:
        return None


def _merge(base: Dict, patch: Optional[Dict]) -> Dict:
    if not patch:
        return base
    out = dict(base)
    for key, value in patch.items():
        if isinstance(value, dict) and isinstance(out.get(key), dict):
            out[key] = _merge(out[key], value)
        else:
            out[key] = value
    return out


def _problem(
    problem_type: str,
    problem_desc: str,
    position: str,
    standard_value: str,
    current_value: str,
    suggestion: str,
    severity: str = "WARNING",
    auto_fix: bool = True,
    confidence: float = 1.0,
    issue_status: str = "open",
    reasons: Optional[List[str]] = None,
    region: Optional[str] = None,
    paragraph_type: Optional[str] = None,
) -> Dict:
    desc = _normalize_problem_desc(problem_desc)
    normalized_suggestion = _normalize_suggestion(suggestion)
    normalized_reasons = _normalize_reasons(reasons or [])
    normalized_position = _normalize_position(position)
    return {
        "problem_type": problem_type,
        "problem_desc": desc,
        "position": normalized_position,
        "standard_value": standard_value,
        "current_value": current_value,
        "suggestion": normalized_suggestion,
        "severity": severity,
        "auto_fix": auto_fix,
        "confidence": round(float(confidence), 2),
        "issue_status": issue_status,
        "reasons": normalized_reasons,
        "region": region,
        "paragraph_type": paragraph_type,
    }


def _normalize_problem_desc(problem_desc: str) -> str:
    text = str(problem_desc or "").strip()
    if not text:
        return text

    # If already Chinese, keep as-is.
    if re.search(r"[\u4e00-\u9fff]", text):
        return text

    mapping = {
        "section count exceeds standard": "分节数量超出规范",
        "paper size mismatch": "纸张大小不符合规范",
        "header missing": "页眉缺失",
        "footer missing": "页脚缺失",
        "page number field not detected": "未检测到页码域",
        "page number position mismatch": "页码位置不符合规范",
        "low confidence heading candidate": "标题识别置信度较低，建议人工确认",
        "heading number not continuous": "标题编号不连续",
        "heading hierarchy jump detected": "标题层级跳跃",
        "body font mismatch": "正文字体不符合规范",
        "body font size mismatch": "正文字号不符合规范",
        "body line spacing mismatch": "正文行距不符合规范",
        "body first line indent mismatch": "正文首行缩进不符合规范",
        "body paragraph spacing mismatch": "正文段前段后间距不符合规范",
        "figure caption may not be below image": "图题可能不在图片下方",
        "table caption may not be above table": "表题可能不在表格上方",
        "image not centered": "图片未居中",
        "toc field or catalog region not detected": "未检测到自动目录域或目录区域",
        "references region exists but contains no entries": "参考文献区域存在但无条目",
        "entry numbering format not recognized": "参考文献条目编号格式无法识别",
        "references region found, but no entries matched supported numbering format": "参考文献区域内未识别到支持的编号格式",
        "multiple numbering formats are mixed in references": "参考文献编号格式混用",
        "duplicate reference number detected": "参考文献编号重复",
        "reference numbers are out of order": "参考文献编号乱序",
        "reference numbering is not continuous": "参考文献编号不连续",
        "reference entry content is empty": "参考文献条目为空",
        "possible references detected near document end, but reference region marker not found": "文末疑似参考文献，但未检测到参考文献区域标记",
        "reference region not detected": "未检测到参考文献区域",
        "reference index not continuous": "参考文献编号不连续",
        "reference entry may be incomplete": "参考文献信息可能不完整",
        "reference type mark missing or invalid": "参考文献类型标识缺失或不规范",
        "cannot resolve effective style value, needs manual review": "无法确定有效格式值，需人工确认",
        "top margin mismatch": "上边距不符合规范",
        "bottom margin mismatch": "下边距不符合规范",
        "left margin mismatch": "左边距不符合规范",
        "right margin mismatch": "右边距不符合规范",
    }
    if text in mapping:
        return mapping[text]

    # Generic pattern handling
    m = re.match(r"^heading (\d) font mismatch$", text)
    if m:
        return f"{m.group(1)}级标题字体不符合规范"
    m = re.match(r"^heading (\d) font size mismatch$", text)
    if m:
        return f"{m.group(1)}级标题字号不符合规范"
    m = re.match(r"^heading (\d) alignment mismatch$", text)
    if m:
        return f"{m.group(1)}级标题对齐方式不符合规范"
    m = re.match(r"^heading (\d) bold mismatch$", text)
    if m:
        return f"{m.group(1)}级标题加粗设置不符合规范"
    m = re.match(r"^(figure|table) index not continuous$", text)
    if m:
        return "图表编号不连续"

    return text


def _normalize_position(position: str) -> str:
    text = str(position or "").strip()
    if not text:
        return text
    if re.search(r"[\u4e00-\u9fff]", text):
        return text

    low = text.lower()
    if low == "document":
        return "文档整体"

    m = re.match(r"^section\s+(\d+)$", low)
    if m:
        return f"第{m.group(1)}节"

    m = re.match(r"^paragraph\s+(\d+)\s*\((.*)\)$", text, flags=re.IGNORECASE)
    if m:
        preview = str(m.group(2) or "").strip()
        return f"第{m.group(1)}段（{preview}）" if preview else f"第{m.group(1)}段"

    m = re.match(r"^paragraph\s+(\d+)$", low)
    if m:
        return f"第{m.group(1)}段"

    mapping = {
        "references region": "参考文献区域",
        "catalog region": "目录区域",
        "unknown": "未知位置",
    }
    return mapping.get(low, text)


def _normalize_suggestion(suggestion: str) -> str:
    text = str(suggestion or "").strip()
    if not text:
        return text
    normalized = text.lower()

    mapping = {
        "check body font manually": "请人工检查正文字体",
        "check body font size manually": "请人工检查正文字号",
        "check body line spacing manually": "请人工检查正文行距",
        "check body first-line indent manually": "请人工检查正文首行缩进",
        "check heading font manually": "请人工检查标题字体",
        "check heading font size manually": "请人工检查标题字号",
        "check heading bold manually": "请人工检查标题加粗设置",
        "check heading alignment manually": "请人工检查标题对齐方式",
        "check heading numbering sequence": "请检查标题编号连续性",
        "verify heading hierarchy levels": "请检查标题层级关系",
        "add reference entries under references heading": "请在参考文献标题下补充有效条目",
        "fill this reference entry content": "请补全该参考文献条目内容",
        "sort or renumber reference entries": "请按顺序调整或重新编号参考文献条目",
        "renumber references continuously": "请将参考文献编号调整为连续",
        "manual review: confirm reference section heading": "需人工确认：请确认参考文献章节标题",
        "manual review: confirm whether references section is present": "需人工确认：请确认文档中是否存在参考文献章节",
        "manual review: verify if this paragraph is a reference entry": "需人工确认：请确认该段是否为参考文献条目",
        "manual review: confirm heading level and style": "需人工确认：请确认标题层级与样式",
        "manual review suggested": "建议人工复核后处理",
        "manual review: use [n], n. or n、 format": "需人工确认：请使用 [n]、n. 或 n、 编号格式",
        "use one consistent numbering format": "请统一参考文献编号格式",
        "set image paragraph alignment to center": "请将图片所在段落设置为居中对齐",
    }
    if normalized in mapping:
        return mapping[normalized]

    def _translate_tail(src: str) -> str:
        out = str(src or "").strip()
        if not out:
            return out
        replacements = [
            (r"top margin", "上边距"),
            (r"bottom margin", "下边距"),
            (r"left margin", "左边距"),
            (r"right margin", "右边距"),
            (r"header content", "页眉内容"),
            (r"footer content", "页脚内容"),
            (r"page number field", "页码域"),
            (r"\bin footer\b", "在页脚中"),
            (r"\balignment\b", "对齐方式"),
            (r"\bcenter\b", "居中"),
            (r"\bleft\b", "左对齐"),
            (r"\bright\b", "右对齐"),
            (r"\bbold\b", "加粗"),
            (r"\btrue\b", "是"),
            (r"\bfalse\b", "否"),
            (r"\bchars?\b", "字符"),
        ]
        for pat, repl in replacements:
            out = re.sub(pat, repl, out, flags=re.IGNORECASE)
        return out.strip()

    m = re.match(r"^set (top|bottom|left|right) margin to ([\d.]+)cm$", text, re.IGNORECASE)
    if m:
        side_map = {"top": "上", "bottom": "下", "left": "左", "right": "右"}
        return f"请将{side_map.get(m.group(1).lower(), m.group(1))}边距设置为 {m.group(2)}cm"

    m = re.match(r"^add header content$", text, re.IGNORECASE)
    if m:
        return "请补充页眉内容"

    m = re.match(r"^add footer content$", text, re.IGNORECASE)
    if m:
        return "请补充页脚内容"

    m = re.match(r"^insert page number field(?: in footer)?$", text, re.IGNORECASE)
    if m:
        return "请在页脚插入页码域"

    m = re.match(r"^insert page number field in footer and align (center|left|right)$", text, re.IGNORECASE)
    if m:
        pos_map = {"center": "居中", "left": "左对齐", "right": "右对齐"}
        return f"请在页脚插入页码域并设置为{pos_map.get(m.group(1).lower(), m.group(1))}"

    m = re.match(r"^use\s+(.+)$", text, re.IGNORECASE)
    if m:
        return f"建议使用{_translate_tail(m.group(1))}"
    m = re.match(r"^set\s+(.+)$", text, re.IGNORECASE)
    if m:
        return f"请设置为{_translate_tail(m.group(1))}"
    m = re.match(r"^check\s+(.+)$", text, re.IGNORECASE)
    if m:
        return f"请检查{_translate_tail(m.group(1))}"
    m = re.match(r"^manual review:\s*(.+)$", text, re.IGNORECASE)
    if m:
        return f"需人工确认：{m.group(1)}"

    m = re.match(r"^set page number alignment to (.+)$", text, re.IGNORECASE)
    if m:
        pos = m.group(1).strip().lower()
        pos_map = {"bottom_center": "页脚居中", "center": "居中", "left": "左对齐", "right": "右对齐"}
        return f"请将页码位置设置为{pos_map.get(pos, pos)}"

    m = re.match(r"^set heading bold to (true|false)$", text, re.IGNORECASE)
    if m:
        return f"请将标题加粗设置为{'加粗' if m.group(1).lower() == 'true' else '不加粗'}"

    m = re.match(r"^use ([\d.]+)pt$", text, re.IGNORECASE)
    if m:
        return f"请使用 {m.group(1)}pt"

    m = re.match(r"^use ([\d.]+) chars$", text, re.IGNORECASE)
    if m:
        return f"请使用 {m.group(1)} 个字符"

    if re.search(r"[\u4e00-\u9fff]", text):
        return text

    return text


def _normalize_reason_item(reason: str) -> str:
    text = str(reason or "").strip()
    if not text:
        return text
    if re.search(r"[\u4e00-\u9fff]", text):
        return text

    mapping = {
        "matched_style_name": "匹配到样式名特征",
        "matched_number_pattern": "匹配到编号规则",
        "short_text": "文本较短，符合标题特征",
        "bold": "检测到加粗特征",
        "larger_font": "字号明显大于正文",
        "slightly_larger_font": "字号略大于正文",
        "center_aligned": "居中对齐特征",
        "left_aligned": "左对齐特征",
        "in_body_region": "位于正文区域",
        "chapter_keyword": "包含章节关键词",
        "borderline_heading_promoted": "边界标题已提升为可确认标题",
        "suspected_heading": "疑似标题",
        "default_body": "默认归类为正文",
        "generic_text": "普通文本特征",
        "no_known_pattern": "未匹配到已知模式",
        "effective_font_name_unknown": "有效字体值无法确定",
        "effective_font_size_unknown": "有效字号值无法确定",
        "effective_bold_unknown": "有效加粗值无法确定",
        "effective_alignment_unknown": "有效对齐值无法确定",
        "effective_line_spacing_unknown": "有效行距值无法确定",
        "effective_first_line_indent_unknown": "有效首行缩进值无法确定",
        "heading font check": "标题字体检查",
        "heading size check": "标题字号检查",
        "heading bold check": "标题加粗检查",
        "heading alignment check": "标题对齐检查",
        "heading sequence check": "标题编号连续性检查",
        "heading hierarchy check": "标题层级检查",
        "body font check": "正文字体检查",
        "body font size check": "正文字号检查",
        "body line spacing check": "正文行距检查",
        "body first line indent check": "正文首行缩进检查",
        "body paragraph spacing check": "正文段间距检查",
        "page margin check": "页边距检查",
        "paper size check": "纸张大小检查",
        "header presence check": "页眉存在性检查",
        "footer presence check": "页脚存在性检查",
        "footer xml page field check": "页码域检查",
        "page number alignment check": "页码位置检查",
        "catalog region / toc field check": "目录区域与目录域检查",
        "reference_region_detected": "检测到参考文献区域",
        "empty_reference_entry": "存在空参考文献条目",
        "unsupported_reference_number_format": "参考文献编号格式不受支持",
        "mixed_reference_numbering_format": "参考文献编号格式混用",
        "duplicate_reference_number": "存在重复参考文献编号",
        "reference_number_order_check": "参考文献编号顺序检查",
        "reference_number_continuity_check": "参考文献编号连续性检查",
        "no_reference_marker": "未检测到参考文献区域标记",
        "numbered_entries_near_end": "文档末尾存在疑似编号条目",
        "no_tail_reference_candidates": "文末未检测到疑似参考文献条目",
        "no_nonempty_entries": "参考文献区域无非空条目",
        "no_supported_numbering_entries": "未识别到支持的参考文献编号格式",
        "figure caption position check": "图题位置检查",
        "table caption position check": "表题位置检查",
        "image alignment check": "图片对齐检查",
        "section count check": "分节数量检查",
    }
    if text in mapping:
        return mapping[text]

    if "_" in text:
        probe = text.replace("_", " ").strip()
        if probe in mapping:
            return mapping[probe]
    return text


def _normalize_reasons(reasons: List[str]) -> List[str]:
    return [_normalize_reason_item(item) for item in reasons]


def _severity(code: str, severities: Dict[str, str], default: str = "WARNING") -> str:
    result = severities.get(code, default)
    return result if result in ("ERROR", "WARNING", "INFO") else default


def _manual_review_problem(
    problem_type: str,
    position: str,
    suggestion: str,
    confidence: float,
    reasons: List[str],
    region: Optional[str] = None,
    paragraph_type: Optional[str] = None,
) -> Dict:
    return _problem(
        problem_type=problem_type,
        problem_desc="cannot resolve effective style value, needs manual review",
        position=position,
        standard_value="known effective value",
        current_value="unknown",
        suggestion=suggestion,
        severity="INFO",
        auto_fix=False,
        confidence=confidence,
        issue_status="manual_review",
        reasons=reasons,
        region=region,
        paragraph_type=paragraph_type,
    )


def _style_name(paragraph) -> str:
    if paragraph.style is None or paragraph.style.name is None:
        return ""
    return str(paragraph.style.name).strip()


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", (text or "")).lower()


def _resolve_document(paragraph, explicit_document=None):
    if explicit_document is not None:
        return explicit_document
    part = getattr(paragraph, "part", None)
    return getattr(part, "document", None) if part is not None else None


def _normal_style(document):
    if document is None:
        return None
    try:
        return document.styles["Normal"]
    except Exception:
        return None


def _rfonts_name(rfonts) -> Optional[str]:
    if rfonts is None:
        return None
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
        try:
            name = rfonts.get(qn(key))
            if name:
                return str(name)
        except Exception:
            continue
    return None


def get_effective_font_name(run, paragraph, document):
    if run is not None:
        try:
            run_rpr = run._element.rPr if run._element is not None else None
            run_name = _rfonts_name(run_rpr.rFonts if run_rpr is not None else None)
            if run_name:
                return run_name
        except Exception:
            pass
        if run.font is not None and run.font.name:
            return run.font.name

    try:
        style_rpr = paragraph.style.element.rPr if paragraph.style is not None else None
        style_rfonts_name = _rfonts_name(style_rpr.rFonts if style_rpr is not None else None)
        if style_rfonts_name:
            return style_rfonts_name
    except Exception:
        pass

    try:
        style_name = paragraph.style.font.name
        if style_name:
            return style_name
    except Exception:
        pass

    normal = _normal_style(document)
    try:
        normal_rpr = normal.element.rPr if normal is not None else None
        normal_rfonts_name = _rfonts_name(normal_rpr.rFonts if normal_rpr is not None else None)
        if normal_rfonts_name:
            return normal_rfonts_name
    except Exception:
        pass

    try:
        default_name = normal.font.name if normal is not None else None
        if default_name:
            return default_name
    except Exception:
        pass
    return "unknown"


def get_effective_font_size(run, paragraph, document):
    if run is not None and run.font is not None and run.font.size is not None:
        size = _pt(run.font.size)
        if size is not None:
            return size

    try:
        style_size = _pt(paragraph.style.font.size)
        if style_size is not None:
            return style_size
    except Exception:
        pass

    normal = _normal_style(document)
    try:
        default_size = _pt(normal.font.size) if normal is not None else None
        if default_size is not None:
            return default_size
    except Exception:
        pass
    return "unknown"


def get_effective_bold(run, paragraph):
    if run is not None and run.bold is not None:
        return bool(run.bold)
    try:
        if paragraph.style.font.bold is not None:
            return bool(paragraph.style.font.bold)
    except Exception:
        pass
    normal = _normal_style(_resolve_document(paragraph))
    try:
        if normal is not None and normal.font.bold is not None:
            return bool(normal.font.bold)
    except Exception:
        pass
    return "unknown"


def get_effective_alignment(paragraph):
    if paragraph.paragraph_format.alignment is not None:
        return paragraph.paragraph_format.alignment
    try:
        if paragraph.style.paragraph_format.alignment is not None:
            return paragraph.style.paragraph_format.alignment
    except Exception:
        pass

    normal = _normal_style(_resolve_document(paragraph))
    try:
        if normal is not None and normal.paragraph_format.alignment is not None:
            return normal.paragraph_format.alignment
    except Exception:
        pass
    return "unknown"


def _line_spacing_value(line_spacing) -> Optional[float]:
    if isinstance(line_spacing, (float, int)):
        return float(line_spacing)
    return None


def get_effective_line_spacing(paragraph):
    value = _line_spacing_value(paragraph.paragraph_format.line_spacing)
    if value is not None:
        return value
    try:
        value = _line_spacing_value(paragraph.style.paragraph_format.line_spacing)
        if value is not None:
            return value
    except Exception:
        pass

    normal = _normal_style(_resolve_document(paragraph))
    try:
        value = _line_spacing_value(normal.paragraph_format.line_spacing) if normal is not None else None
        if value is not None:
            return value
    except Exception:
        pass
    return "unknown"


def get_effective_first_line_indent(paragraph):
    indent = _pt(paragraph.paragraph_format.first_line_indent)
    if indent is not None:
        return indent
    try:
        indent = _pt(paragraph.style.paragraph_format.first_line_indent)
        if indent is not None:
            return indent
    except Exception:
        pass

    normal = _normal_style(_resolve_document(paragraph))
    try:
        indent = _pt(normal.paragraph_format.first_line_indent) if normal is not None else None
        if indent is not None:
            return indent
    except Exception:
        pass
    return "unknown"


def _get_paragraph_font_name(paragraph, document=None) -> Optional[str]:
    doc = _resolve_document(paragraph, document)
    for run in paragraph.runs:
        if run.text and run.text.strip():
            name = get_effective_font_name(run, paragraph, doc)
            if name != "unknown":
                return name
    return get_effective_font_name(None, paragraph, doc)


def _get_paragraph_font_size_pt(paragraph, document=None) -> Optional[float]:
    doc = _resolve_document(paragraph, document)
    for run in paragraph.runs:
        if run.text and run.text.strip():
            size = get_effective_font_size(run, paragraph, doc)
            if size != "unknown":
                return float(size)
    size = get_effective_font_size(None, paragraph, doc)
    return None if size == "unknown" else float(size)


def _heading_number_level(text: str) -> Optional[int]:
    text = (text or "").strip()
    if not text:
        return None

    lv1_patterns = (
        rf"^\u7b2c[\u4e00-\u9fa5\d]+{KW_CHAPTER}\s*",
        rf"^[{CN_NUM}]+\u3001\s*",
        r"^\d+\.(?!\d)\s*",
    )
    lv2_patterns = (
        r"^\d+\.\d+(?:\.|\u3001)?\s*",
        rf"^\uff08[{CN_NUM}]+\uff09\s*",
        rf"^\([{CN_NUM}]+\)\s*",
    )
    lv3_patterns = (
        r"^\d+\.\d+\.\d+(?:\.|\u3001)?\s*",
    )
    for pat in lv3_patterns:
        if re.match(pat, text):
            return 3
    for pat in lv2_patterns:
        if re.match(pat, text):
            return 2
    for pat in lv1_patterns:
        if re.match(pat, text):
            return 1
    return None


def _parse_heading_numbers(text: str) -> Optional[Tuple[int, ...]]:
    text = (text or "").strip()
    match = re.match(r"^(\d+(?:\.\d+){0,2})(?:\.)?(?=\D|$)", text)
    if not match:
        return None
    try:
        return tuple(int(x) for x in match.group(1).split("."))
    except Exception:
        return None


def _looks_like_toc_entry_text(text: str) -> bool:
    """
    Heuristic detection for TOC lines, e.g.
    - 第一章 绪论\t1
    - 1.2.3 标题 ...... 12
    - 2.1 标题 8
    """
    raw = (text or "").strip()
    if not raw:
        return False
    compact = re.sub(r"\s+", " ", raw)

    if re.search(r"\t\s*\d+\s*$", raw):
        return True
    if re.search(r"(?:\.{2,}|·{3,}|…{2,}|-{3,}|_{3,}|—{2,}|－{2,})\s*\d+\s*$", compact):
        return True
    if re.match(r"^(第[一二三四五六七八九十百千万0-9]+[章节]|[0-9]+(?:\.[0-9]+){0,3}\.?)\s*.+\s+\d+\s*$", compact):
        return True
    return False


def _is_heading_style(paragraph, level: int) -> bool:
    style = _style_name(paragraph).lower()
    if level == 1 and ("heading 1" in style or f"{CN_TITLE} 1" in style or f"{CN_TITLE}1" in style):
        return True
    if level == 2 and ("heading 2" in style or f"{CN_TITLE} 2" in style or f"{CN_TITLE}2" in style):
        return True
    if level == 3 and ("heading 3" in style or f"{CN_TITLE} 3" in style or f"{CN_TITLE}3" in style):
        return True
    return False


def _heading_level(paragraph) -> Optional[int]:
    for level in (1, 2, 3):
        if _is_heading_style(paragraph, level):
            return level
    return _heading_number_level(paragraph.text or "")


def _heading_style_level(paragraph) -> Optional[int]:
    for level in (1, 2, 3):
        if _is_heading_style(paragraph, level):
            return level
    return None


def _is_bold_paragraph(paragraph) -> bool:
    text_runs = [run for run in paragraph.runs if run.text and run.text.strip()]
    if not text_runs:
        return False
    bold_count = 0
    known_count = 0
    for run in text_runs:
        value = get_effective_bold(run, paragraph)
        if value == "unknown":
            continue
        known_count += 1
        if value:
            bold_count += 1
    return known_count > 0 and (bold_count / known_count >= 0.5)


def _paragraph_font_size_stats(paragraph) -> Tuple[Optional[float], Optional[float]]:
    doc = _resolve_document(paragraph)
    sizes = [
        get_effective_font_size(run, paragraph, doc)
        for run in paragraph.runs
        if run.text and run.text.strip()
    ]
    sizes = [x for x in sizes if x is not None]
    sizes = [float(x) for x in sizes if x != "unknown"]
    if not sizes:
        return None, None
    return sum(sizes) / len(sizes), max(sizes)


def _contains_chapter_keyword(text: str) -> bool:
    value = (text or "").strip()
    return any(keyword in value for keyword in CN_CHAPTER_KEYWORDS)


def _is_sentence_like(text: str) -> bool:
    value = (text or "").strip()
    return bool(value) and value[-1] in {".", "!", "?", "\u3002", "\uff01", "\uff1f", ";", "\uff1b"}


def _is_body_region(region: str) -> bool:
    return region in {"body", "appendix"}


def _marker_region(text: str) -> Optional[str]:
    raw = (text or "").strip()
    if not raw:
        return None
    value = _normalize_text(raw)

    if value in {KW_ABSTRACT, "abstract", KW_CN_ABSTRACT, KW_EN_ABSTRACT}:
        return "abstract"
    if value in {KW_CATALOG, "contents", "tableofcontents"}:
        return "catalog"
    if value in {KW_REFERENCES, "references"}:
        return "references"
    if value in {KW_THANKS, KW_ACK, "acknowledgements", "acknowledgments"}:
        return "thanks"

    # Robust markers for OCR/export cases where heading is followed by page number punctuation.
    if re.match(r"^(摘要|中文摘要|英文摘要)([\s:：·\-\.\d（）\(\)]*)$", raw, flags=re.IGNORECASE):
        return "abstract"
    if re.match(r"^(目录|contents|table\s*of\s*contents)([\s:：·\-\.\d（）\(\)]*)$", raw, flags=re.IGNORECASE):
        return "catalog"
    if re.match(r"^(参考文献|references)([\s:：·\-\.\d（）\(\)]*)$", raw, flags=re.IGNORECASE):
        return "references"
    if re.match(r"^(致谢|谢辞|acknowledgements?|结束语|结论)([\s:：·\-\.\d（）\(\)]*)$", raw, flags=re.IGNORECASE):
        return "thanks"

    if value.startswith(KW_APPENDIX) or value.startswith("appendix"):
        return "appendix"
    return None


def _is_reference_item(text: str) -> bool:
    value = (text or "").strip()
    if re.match(r"^\[\d+\]", value):
        return True
    if re.match(r"^\d+\.", value):
        return True
    if re.match(r"^\d+\u3001", value):
        return True
    if len(value) >= 12 and any(token in value for token in ("[J]", "[M]", "[C]", "[D]", "[R]", "[N]", "[P]", "doi")):
        return True
    return False


def _is_figure_caption(text: str) -> bool:
    pattern = rf"^({KW_FIGURE}|figure)\s*[\d\-\.]+"
    return bool(re.match(pattern, (text or "").strip(), flags=re.IGNORECASE))


def _is_table_caption(text: str) -> bool:
    pattern = rf"^({KW_TABLE}|table)\s*[\d\-\.]+"
    return bool(re.match(pattern, (text or "").strip(), flags=re.IGNORECASE))


def _infer_body_start(paragraphs) -> Optional[int]:
    for idx, paragraph in enumerate(paragraphs, start=1):
        text = (paragraph.text or "").strip()
        if not text:
            continue
        if _marker_region(text):
            continue
        if _is_figure_caption(text) or _is_table_caption(text) or _is_reference_item(text):
            continue
        if _heading_level(paragraph) == 1:
            return idx

    for idx, paragraph in enumerate(paragraphs, start=1):
        text = (paragraph.text or "").strip()
        if len(text) >= 20 and not _marker_region(text):
            return idx
    return None


def _next_marker_after(pos: int, points: List[Optional[int]]) -> Optional[int]:
    values = [x for x in points if x is not None and x > pos]
    return min(values) if values else None


def _section_paragraph_ranges(doc: Document) -> List[Tuple[int, int]]:
    paragraphs = doc.paragraphs
    if not paragraphs:
        return []
    boundaries = []
    for idx, paragraph in enumerate(paragraphs, start=1):
        try:
            ppr = paragraph._p.pPr
            if ppr is not None and ppr.sectPr is not None:
                boundaries.append(idx)
        except Exception:
            continue
    ranges: List[Tuple[int, int]] = []
    start = 1
    for end in boundaries:
        if end >= start:
            ranges.append((start, end))
            start = end + 1
    if start <= len(paragraphs):
        ranges.append((start, len(paragraphs)))
    if not ranges:
        ranges.append((1, len(paragraphs)))
    return ranges


def _parse_exclude_regions(value) -> Set[str]:
    if value is None:
        return set()
    if isinstance(value, (list, tuple, set)):
        raw_items = [str(item).strip() for item in value]
    else:
        raw_items = [x.strip() for x in str(value).split(",")]
    mapping = {
        "cover": "cover",
        "abstract": "abstract",
        "abstract_cn": "abstract",
        "abstract_en": "abstract",
        "toc": "catalog",
        "catalog": "catalog",
        "contents": "catalog",
        "references": "references",
        "thanks": "thanks",
        "ack": "thanks",
        "appendix": "appendix",
        "body": "body",
        "unknown": "unknown",
    }
    out = set()
    for item in raw_items:
        if not item:
            continue
        key = item.strip().lower()
        out.add(mapping.get(key, key))
    return out


def classify_paragraph(paragraph, context: Dict) -> Dict:
    text = (paragraph.text or "").strip()
    region = context.get("region", "unknown")
    body_font_size = float(context.get("body_font_size_pt", 12.0))

    if not text:
        return {"type": "empty", "confidence": 1.0, "reasons": ["empty_paragraph"], "issue_status": "open"}

    marker = _marker_region(text)
    if marker:
        return {
            "type": "section_marker",
            "confidence": 0.99,
            "reasons": ["region_marker"],
            "issue_status": "open",
        }

    if _is_figure_caption(text):
        return {"type": "figure_caption", "confidence": 0.96, "reasons": ["figure_caption_pattern"], "issue_status": "open"}

    if _is_table_caption(text):
        return {"type": "table_caption", "confidence": 0.96, "reasons": ["table_caption_pattern"], "issue_status": "open"}

    if region == "references" and _is_reference_item(text):
        return {"type": "reference_item", "confidence": 0.92, "reasons": ["reference_item_pattern"], "issue_status": "open"}

    style_level = _heading_style_level(paragraph)
    number_level = _heading_number_level(text)
    is_bold = _is_bold_paragraph(paragraph)
    avg_size, max_size = _paragraph_font_size_stats(paragraph)
    align = get_effective_alignment(paragraph)
    text_len = len(text)
    has_keyword = _contains_chapter_keyword(text)
    in_body = _is_body_region(region)
    sentence_like = _is_sentence_like(text)
    explicit_heading_signal = style_level in (1, 2, 3) or number_level in (1, 2, 3)
    strong_visual_heading = bool(
        is_bold
        and max_size is not None
        and max_size >= body_font_size + 1
        and text_len <= 30
        and not sentence_like
    )
    compact_numbered_title = bool(
        number_level in (1, 2, 3)
        and text_len <= 26
        and not sentence_like
        and not re.search(r"[，,。；;：:]", text)
    )

    scores = {1: 0.0, 2: 0.0, 3: 0.0}
    reasons = {1: [], 2: [], 3: []}

    if style_level in (1, 2, 3):
        scores[style_level] += 0.32
        reasons[style_level].append("matched_style_name")

    if number_level in (1, 2, 3):
        scores[number_level] += 0.40
        reasons[number_level].append("matched_number_pattern")
        if compact_numbered_title:
            scores[number_level] += 0.14
            reasons[number_level].append("compact_numbered_heading")

    if text_len <= 30:
        for lv in scores:
            scores[lv] += 0.08
            reasons[lv].append("short_text")
    elif text_len >= 48:
        for lv in scores:
            scores[lv] -= 0.12

    if is_bold:
        for lv in scores:
            scores[lv] += 0.1
            reasons[lv].append("bold")

    if max_size is not None:
        if max_size >= body_font_size + 2:
            scores[1] += 0.14
            scores[2] += 0.08
            reasons[1].append("larger_font")
            reasons[2].append("larger_font")
        elif max_size >= body_font_size + 1:
            for lv in scores:
                scores[lv] += 0.05
                reasons[lv].append("slightly_larger_font")

    if align == WD_ALIGN_PARAGRAPH.CENTER:
        scores[1] += 0.08
        reasons[1].append("center_aligned")
    elif align in (WD_ALIGN_PARAGRAPH.LEFT, None, "unknown"):
        scores[2] += 0.05
        scores[3] += 0.05
        reasons[2].append("left_aligned")
        reasons[3].append("left_aligned")

    if in_body:
        for lv in scores:
            scores[lv] += 0.08
            reasons[lv].append("in_body_region")
    else:
        for lv in scores:
            scores[lv] -= 0.14

    if has_keyword:
        scores[1] += 0.08
        scores[2] += 0.05
        reasons[1].append("chapter_keyword")
        reasons[2].append("chapter_keyword")

    if sentence_like:
        for lv in scores:
            scores[lv] -= 0.08

    best_level = max(scores, key=lambda lv: scores[lv])
    best_score = max(0.0, min(1.0, scores[best_level]))
    sorted_scores = sorted(scores.values(), reverse=True)
    margin = sorted_scores[0] - sorted_scores[1]

    if best_score >= HEADING_HIGH_CONFIDENCE and margin >= 0.05 and (explicit_heading_signal or strong_visual_heading):
        return {
            "type": f"heading_{best_level}",
            "confidence": round(best_score, 2),
            "reasons": sorted(set(reasons[best_level])),
            "issue_status": "open",
        }

    # Promote borderline heading candidates when explicit numbering/style signal exists.
    # This increases one-click format coverage while keeping manual review for weak signals.
    if best_score >= 0.74 and margin >= 0.04 and in_body and explicit_heading_signal:
        promoted_reasons = sorted(set(reasons[best_level] + ["borderline_heading_promoted"]))
        return {
            "type": f"heading_{best_level}",
            "confidence": round(best_score, 2),
            "reasons": promoted_reasons,
            "issue_status": "open",
        }

    if best_score >= HEADING_LOW_CONFIDENCE and margin >= 0.02 and (explicit_heading_signal or strong_visual_heading):
        return {
            "type": "body",
            "confidence": round(best_score, 2),
            "reasons": sorted(set(reasons[best_level])) + ["suspected_heading"],
            "issue_status": "manual_review",
            "suspected_type": f"heading_{best_level}",
        }

    if in_body:
        return {"type": "body", "confidence": 0.84, "reasons": ["default_body", "in_body_region"], "issue_status": "open"}
    if text_len >= 8:
        return {"type": "body", "confidence": 0.68, "reasons": ["generic_text"], "issue_status": "open"}
    return {"type": "unknown", "confidence": 0.5, "reasons": ["no_known_pattern"], "issue_status": "open"}


def _analyze_document_structure(doc: Document, standards: Optional[Dict] = None) -> List[Dict]:
    paragraphs = doc.paragraphs
    body_font_size = float((standards or DEFAULT_STANDARDS).get("body", {}).get("font_size_pt", 12.0))
    index_map: Dict[str, Optional[int]] = {
        "abstract": None,
        "catalog": None,
        "references": None,
        "thanks": None,
        "appendix": None,
    }

    for idx, paragraph in enumerate(paragraphs, start=1):
        marker = _marker_region(paragraph.text or "")
        if marker in index_map and index_map[marker] is None:
            index_map[marker] = idx

    body_start = _infer_body_start(paragraphs)
    region_by_index = {idx: "unknown" for idx in range(1, len(paragraphs) + 1)}

    anchors = [
        index_map["abstract"],
        index_map["catalog"],
        body_start,
        index_map["references"],
        index_map["thanks"],
        index_map["appendix"],
    ]
    first_anchor = min([x for x in anchors if x is not None], default=None)
    if first_anchor is not None:
        for idx in range(1, first_anchor):
            region_by_index[idx] = "cover"

    if index_map["abstract"] is not None:
        end = _next_marker_after(
            index_map["abstract"],
            [index_map["catalog"], body_start, index_map["references"], index_map["thanks"], index_map["appendix"]],
        )
        for idx in range(index_map["abstract"], (end or len(paragraphs) + 1)):
            region_by_index[idx] = "abstract"

    if index_map["catalog"] is not None:
        end = _next_marker_after(index_map["catalog"], [body_start, index_map["references"], index_map["thanks"], index_map["appendix"]])
        for idx in range(index_map["catalog"], (end or len(paragraphs) + 1)):
            region_by_index[idx] = "catalog"

    if body_start is not None:
        end = _next_marker_after(body_start, [index_map["references"], index_map["thanks"], index_map["appendix"]])
        for idx in range(body_start, (end or len(paragraphs) + 1)):
            # Keep catalog region intact; TOC entries must not be re-labeled as body.
            if region_by_index[idx] in {"unknown", "cover", "abstract"}:
                region_by_index[idx] = "body"

    if index_map["references"] is not None:
        end = _next_marker_after(index_map["references"], [index_map["thanks"], index_map["appendix"]])
        for idx in range(index_map["references"], (end or len(paragraphs) + 1)):
            region_by_index[idx] = "references"

    if index_map["thanks"] is not None:
        end = _next_marker_after(index_map["thanks"], [index_map["appendix"]])
        for idx in range(index_map["thanks"], (end or len(paragraphs) + 1)):
            region_by_index[idx] = "thanks"

    if index_map["appendix"] is not None:
        for idx in range(index_map["appendix"], len(paragraphs) + 1):
            region_by_index[idx] = "appendix"

    contexts: List[Dict] = []
    for idx, paragraph in enumerate(paragraphs, start=1):
        region = region_by_index[idx] if region_by_index[idx] in REGIONS else "unknown"
        classify_result = classify_paragraph(
            paragraph,
            {
                "region": region,
                "index": idx,
                "body_font_size_pt": body_font_size,
            },
        )
        paragraph_type = classify_result.get("type", "unknown")
        confidence = float(classify_result.get("confidence", 0.5))
        reasons = list(classify_result.get("reasons", []))
        issue_status = classify_result.get("issue_status", "open")
        suspected_type = classify_result.get("suspected_type")
        if paragraph_type not in PARAGRAPH_TYPES:
            paragraph_type = "unknown"
        contexts.append(
            {
                "index": idx,
                "paragraph": paragraph,
                "text": (paragraph.text or "").strip(),
                "region": region,
                "paragraph_type": paragraph_type,
                "confidence": confidence,
                "reasons": reasons,
                "issue_status": issue_status,
                "suspected_type": suspected_type,
            }
        )
    return contexts


def _detect_page(
    doc: Document,
    standards: Dict,
    checks: Dict[str, bool],
    severities: Dict[str, str],
    contexts: Optional[List[Dict]] = None,
) -> List[Dict]:
    problems = []
    margin = standards["page_margin_cm"]
    page = standards["page"]
    max_sections = page.get("max_sections", 1)
    if max_sections is not None:
        try:
            max_sections = int(max_sections)
        except Exception:
            max_sections = 1

    header_start_from_body = bool(page.get("header_footer_start_from_body", False))
    page_number_start_from_body = bool(page.get("page_number_start_from_body", False))
    header_exclude_regions = _parse_exclude_regions(page.get("header_footer_exclude_regions"))
    page_number_exclude_regions = _parse_exclude_regions(page.get("page_number_exclude_regions"))
    body_like_regions = {"body", "appendix", "references"}

    region_by_paragraph = {}
    if contexts:
        region_by_paragraph = {int(ctx.get("index", 0)): str(ctx.get("region", "unknown")) for ctx in contexts}
    section_ranges = _section_paragraph_ranges(doc)

    if checks.get("SECTION_BREAK", True) and max_sections is not None and len(doc.sections) > max_sections:
        problems.append(
            _problem(
                "section_break",
                "section count exceeds standard",
                "document",
                f"sections <= {max_sections}",
                f"{len(doc.sections)}",
                "remove unnecessary section breaks",
                severity=_severity("SECTION_BREAK", severities, "WARNING"),
                auto_fix=False,
                confidence=0.98,
                reasons=["section count check"],
            )
        )

    for sec_idx, section in enumerate(doc.sections, start=1):
        pos = f"section {sec_idx}"
        if sec_idx - 1 < len(section_ranges):
            start_idx, end_idx = section_ranges[sec_idx - 1]
        else:
            start_idx, end_idx = 1, len(doc.paragraphs)

        section_regions = [region_by_paragraph.get(i, "unknown") for i in range(start_idx, end_idx + 1)]
        known_regions = [x for x in section_regions if x and x != "unknown"]
        if known_regions:
            region_counts: Dict[str, int] = {}
            for item in known_regions:
                region_counts[item] = region_counts.get(item, 0) + 1
            section_region = max(region_counts.items(), key=lambda kv: kv[1])[0]
        else:
            section_region = "unknown"

        if checks.get("PAGE_MARGIN", True):
            values = {
                "top": _cm(section.top_margin),
                "bottom": _cm(section.bottom_margin),
                "left": _cm(section.left_margin),
                "right": _cm(section.right_margin),
            }
            for side in ("top", "bottom", "left", "right"):
                if values[side] is None:
                    continue
                if abs(values[side] - float(margin[side])) > 0.05:
                    problems.append(
                        _problem(
                            "page_margin",
                            f"{side} margin mismatch",
                            pos,
                            f"{margin[side]}cm",
                            f"{values[side]:.2f}cm",
                            f"set {side} margin to {margin[side]}cm",
                            severity=_severity("PAGE_MARGIN", severities, "ERROR"),
                            auto_fix=True,
                            confidence=0.95,
                            reasons=["page margin check"],
                        )
                    )

        if checks.get("PAGE_SIZE", True):
            width = _cm(section.page_width)
            height = _cm(section.page_height)
            std_w = float(page.get("paper_width_cm", 21.0))
            std_h = float(page.get("paper_height_cm", 29.7))
            if width is not None and height is not None:
                ok = (abs(width - std_w) <= 0.1 and abs(height - std_h) <= 0.1) or (abs(width - std_h) <= 0.1 and abs(height - std_w) <= 0.1)
                if not ok:
                    problems.append(
                        _problem(
                            "paper_size",
                            "paper size mismatch",
                            pos,
                            f"{std_w}cm x {std_h}cm",
                            f"{width:.2f}cm x {height:.2f}cm",
                            "set paper size to A4",
                            severity=_severity("PAGE_SIZE", severities, "ERROR"),
                            auto_fix=True,
                            confidence=0.95,
                            reasons=["paper size check"],
                        )
                    )

        if checks.get("HEADER_FOOTER", True):
            if section_region in header_exclude_regions:
                continue
            if header_start_from_body and section_region not in body_like_regions:
                continue
            header_text = section.header.paragraphs[0].text.strip() if section.header.paragraphs else ""
            footer_text = section.footer.paragraphs[0].text.strip() if section.footer.paragraphs else ""
            if page.get("require_header") and not header_text:
                problems.append(
                    _problem(
                        "header_footer",
                        "header missing",
                        pos,
                        "header required",
                        "missing",
                        "add header content",
                        severity=_severity("HEADER_FOOTER", severities, "INFO"),
                        auto_fix=False,
                        confidence=0.9,
                        reasons=["header presence check"],
                    )
                )
            if page.get("require_footer") and not footer_text:
                problems.append(
                    _problem(
                        "header_footer",
                        "footer missing",
                        pos,
                        "footer required",
                        "missing",
                        "add footer content",
                        severity=_severity("HEADER_FOOTER", severities, "INFO"),
                        auto_fix=False,
                        confidence=0.9,
                        reasons=["footer presence check"],
                    )
                )

        if checks.get("PAGE_NUMBER", True):
            if section_region in page_number_exclude_regions:
                continue
            if page_number_start_from_body and section_region not in body_like_regions:
                continue
            footer_xml = section.footer._element.xml
            has_page_field = "PAGE" in footer_xml
            target_pos = page.get("page_number_position", "bottom_center")
            align = section.footer.paragraphs[0].paragraph_format.alignment if section.footer.paragraphs else None
            align_map = {
                "bottom_left": WD_ALIGN_PARAGRAPH.LEFT,
                "bottom_center": WD_ALIGN_PARAGRAPH.CENTER,
                "bottom_right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            if not has_page_field:
                problems.append(
                    _problem(
                        "page_number",
                        "page number field not detected",
                        pos,
                        "PAGE field exists",
                        "not found",
                        "insert page number field in footer",
                        severity=_severity("PAGE_NUMBER", severities, "INFO"),
                        auto_fix=False,
                        confidence=0.92,
                        reasons=["footer xml page field check"],
                    )
                )
            elif align is not None and align_map.get(target_pos) is not None and align != align_map[target_pos]:
                problems.append(
                    _problem(
                        "page_number",
                        "page number position mismatch",
                        pos,
                        target_pos,
                        str(align),
                        f"set page number alignment to {target_pos}",
                        severity=_severity("PAGE_NUMBER", severities, "INFO"),
                        auto_fix=False,
                        confidence=0.88,
                        reasons=["page number alignment check"],
                    )
                )

    return problems


def _detect_headings(contexts: List[Dict], standards: Dict, checks: Dict[str, bool], severities: Dict[str, str]) -> List[Dict]:
    problems = []
    heading_info = []
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}

    for ctx in contexts:
        # Do not treat TOC entries as body headings, even when region tagging is off.
        if _looks_like_toc_entry_text(ctx.get("text", "")):
            continue

        if ctx.get("issue_status") in {"suspected", "manual_review"} and ctx.get("suspected_type"):
            problems.append(
                _problem(
                    "suspected_heading",
                    "low confidence heading candidate",
                    f"paragraph {ctx['index']} ({ctx['text'][:20]})",
                    "high-confidence heading required",
                    ctx.get("suspected_type", "heading_unknown"),
                    "manual review suggested",
                    severity="INFO",
                    auto_fix=False,
                    confidence=ctx["confidence"],
                    issue_status="manual_review",
                    reasons=ctx["reasons"],
                    region=ctx["region"],
                    paragraph_type=ctx.get("suspected_type"),
                )
            )
            continue

        if ctx["region"] not in {"body", "appendix"}:
            continue
        if ctx["paragraph_type"] not in {"heading_1", "heading_2", "heading_3"}:
            continue

        paragraph = ctx["paragraph"]
        text = ctx["text"]
        level = int(ctx["paragraph_type"].split("_")[-1])
        heading_info.append((ctx["index"], level, text))

        code = f"HEADING_LEVEL_{level}"
        if not checks.get(code, True):
            continue

        expect = standards["headings"][f"h{level}"]
        pos = f"paragraph {ctx['index']} ({text[:20]})"

        font_name = _get_paragraph_font_name(paragraph)
        if font_name in (None, "unknown"):
            problems.append(
                _manual_review_problem(
                    "heading_font_unknown",
                    pos,
                    "check heading font manually",
                    ctx["confidence"],
                    ctx["reasons"] + ["effective_font_name_unknown"],
                    region=ctx["region"],
                    paragraph_type=ctx["paragraph_type"],
                )
            )
        elif font_name != expect["font_name"]:
            problems.append(
                _problem(
                    "heading_format",
                    f"heading {level} font mismatch",
                    pos,
                    expect["font_name"],
                    font_name,
                    f"use {expect['font_name']}",
                    severity=_severity(code, severities, "WARNING"),
                    auto_fix=True,
                    confidence=ctx["confidence"],
                    reasons=ctx["reasons"] + ["heading font check"],
                    region=ctx["region"],
                    paragraph_type=ctx["paragraph_type"],
                )
            )

        font_size = _get_paragraph_font_size_pt(paragraph)
        if font_size is None:
            problems.append(
                _manual_review_problem(
                    "heading_font_size_unknown",
                    pos,
                    "check heading font size manually",
                    ctx["confidence"],
                    ctx["reasons"] + ["effective_font_size_unknown"],
                    region=ctx["region"],
                    paragraph_type=ctx["paragraph_type"],
                )
            )
        elif abs(font_size - float(expect["font_size_pt"])) > 0.1:
            problems.append(
                _problem(
                    "heading_format",
                    f"heading {level} font size mismatch",
                    pos,
                    f"{expect['font_size_pt']}pt",
                    f"{font_size:.1f}pt",
                    f"use {expect['font_size_pt']}pt",
                    severity=_severity(code, severities, "WARNING"),
                    auto_fix=True,
                    confidence=ctx["confidence"],
                    reasons=ctx["reasons"] + ["heading size check"],
                    region=ctx["region"],
                    paragraph_type=ctx["paragraph_type"],
                )
            )

        text_runs = [run for run in paragraph.runs if run.text and run.text.strip()]
        if text_runs:
            bold_values = [get_effective_bold(run, paragraph) for run in text_runs]
            known_bold = [bool(v) for v in bold_values if v != "unknown"]
            if not known_bold:
                problems.append(
                    _manual_review_problem(
                        "heading_bold_unknown",
                        pos,
                        "check heading bold manually",
                        ctx["confidence"],
                        ctx["reasons"] + ["effective_bold_unknown"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )
            else:
                effective_bold = sum(1 for v in known_bold if v) / len(known_bold) >= 0.5
                expected_bold = bool(expect.get("bold", True))
                if effective_bold != expected_bold:
                    problems.append(
                        _problem(
                            "heading_format",
                            f"heading {level} bold mismatch",
                            pos,
                            str(expected_bold),
                            str(effective_bold),
                            f"set heading bold to {expected_bold}",
                            severity=_severity(code, severities, "INFO"),
                            auto_fix=True,
                            confidence=ctx["confidence"],
                            reasons=ctx["reasons"] + ["heading bold check"],
                            region=ctx["region"],
                            paragraph_type=ctx["paragraph_type"],
                        )
                    )

        expected_align = align_map.get(expect.get("align", "left"))
        effective_align = get_effective_alignment(paragraph)
        if effective_align == "unknown":
            problems.append(
                _manual_review_problem(
                    "heading_alignment_unknown",
                    pos,
                    "check heading alignment manually",
                    ctx["confidence"],
                    ctx["reasons"] + ["effective_alignment_unknown"],
                    region=ctx["region"],
                    paragraph_type=ctx["paragraph_type"],
                )
            )
        elif expected_align is not None and effective_align != expected_align:
            problems.append(
                _problem(
                    "heading_format",
                    f"heading {level} alignment mismatch",
                    pos,
                    expect.get("align", "left"),
                    str(effective_align),
                    f"use {expect.get('align', 'left')}",
                    severity=_severity(code, severities, "INFO"),
                    auto_fix=True,
                    confidence=ctx["confidence"],
                    reasons=ctx["reasons"] + ["heading alignment check"],
                    region=ctx["region"],
                    paragraph_type=ctx["paragraph_type"],
                )
            )

    if checks.get("HEADING_NUMBER", True):
        last_by_level = {1: None, 2: None, 3: None}
        prev_level = 0
        for idx, level, text in heading_info:
            numbers = _parse_heading_numbers(text)
            if numbers:
                current = numbers[-1]
                if last_by_level[level] is not None and current != last_by_level[level] + 1:
                    problems.append(
                        _problem(
                            "heading_number",
                            "heading number not continuous",
                            f"paragraph {idx} ({text[:20]})",
                            str(last_by_level[level] + 1),
                            str(current),
                            "fix heading number continuity",
                            severity=_severity("HEADING_NUMBER", severities, "ERROR"),
                            auto_fix=False,
                            confidence=0.88,
                            reasons=["heading sequence check"],
                            region="body",
                            paragraph_type=f"heading_{level}",
                        )
                    )
                last_by_level[level] = current

            if prev_level and level - prev_level > 1:
                problems.append(
                    _problem(
                        "heading_hierarchy",
                        "heading hierarchy jump detected",
                        f"paragraph {idx} ({text[:20]})",
                        f"max jump: {prev_level} -> {prev_level + 1}",
                        f"{prev_level} -> {level}",
                        "fix heading hierarchy",
                        severity=_severity("HEADING_NUMBER", severities, "ERROR"),
                        auto_fix=False,
                        confidence=0.86,
                        reasons=["heading hierarchy check"],
                        region="body",
                        paragraph_type=f"heading_{level}",
                    )
                )
            prev_level = level

    return problems


def _detect_section_titles(contexts: List[Dict], standards: Dict, checks: Dict[str, bool], severities: Dict[str, str]) -> List[Dict]:
    problems: List[Dict] = []
    title_cfg = standards.get("section_titles", {}) or {}
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}

    def _pick_context(region: str, keyword_fn):
        for ctx in contexts:
            if ctx.get("paragraph_type") != "section_marker":
                continue
            if ctx.get("region") != region:
                continue
            txt = str(ctx.get("text") or "").strip()
            if keyword_fn(txt):
                return ctx
        return None

    def _check(code: str, label: str, ctx: Optional[Dict], cfg_key: str):
        if not checks.get(code, True):
            return
        expect = title_cfg.get(cfg_key) or {}
        if not expect:
            return
        if not ctx:
            return

        paragraph = ctx["paragraph"]
        text = ctx.get("text") or ""
        pos = f"paragraph {ctx['index']} ({text[:20]})"

        expected_font = expect.get("font_name")
        if expected_font:
            font_name = _get_paragraph_font_name(paragraph)
            if font_name not in (None, "unknown") and font_name != expected_font:
                problems.append(
                    _problem(
                        "section_title_format",
                        f"{label} title font mismatch",
                        pos,
                        expected_font,
                        str(font_name),
                        f"use {expected_font}",
                        severity=_severity(code, severities, "ERROR"),
                        auto_fix=True,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + [f"{label} title font check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

        expected_size = expect.get("font_size_pt")
        if expected_size is not None:
            font_size = _get_paragraph_font_size_pt(paragraph)
            if font_size is not None and abs(float(font_size) - float(expected_size)) > 0.1:
                problems.append(
                    _problem(
                        "section_title_format",
                        f"{label} title font size mismatch",
                        pos,
                        f"{expected_size}pt",
                        f"{font_size:.1f}pt",
                        f"use {expected_size}pt",
                        severity=_severity(code, severities, "ERROR"),
                        auto_fix=True,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + [f"{label} title size check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

        expected_bold = expect.get("bold")
        if expected_bold is not None:
            text_runs = [run for run in paragraph.runs if run.text and run.text.strip()]
            known_bold = [get_effective_bold(run, paragraph) for run in text_runs]
            known_bold = [bool(v) for v in known_bold if v != "unknown"]
            if known_bold:
                effective_bold = sum(1 for v in known_bold if v) / len(known_bold) >= 0.5
                if bool(effective_bold) != bool(expected_bold):
                    problems.append(
                        _problem(
                            "section_title_format",
                            f"{label} title bold mismatch",
                            pos,
                            str(bool(expected_bold)),
                            str(bool(effective_bold)),
                            f"set {label} title bold to {bool(expected_bold)}",
                            severity=_severity(code, severities, "WARNING"),
                            auto_fix=True,
                            confidence=ctx["confidence"],
                            reasons=ctx["reasons"] + [f"{label} title bold check"],
                            region=ctx["region"],
                            paragraph_type=ctx["paragraph_type"],
                        )
                    )

        expected_align_name = str(expect.get("align", "") or "").strip().lower()
        expected_align = align_map.get(expected_align_name)
        if expected_align is not None:
            effective_align = get_effective_alignment(paragraph)
            if effective_align != "unknown" and effective_align != expected_align:
                problems.append(
                    _problem(
                        "section_title_format",
                        f"{label} title alignment mismatch",
                        pos,
                        expected_align_name,
                        str(effective_align),
                        f"use {expected_align_name}",
                        severity=_severity(code, severities, "WARNING"),
                        auto_fix=True,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + [f"{label} title alignment check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

    zh_abs_ctx = _pick_context("abstract", lambda t: (KW_ABSTRACT in t or KW_CN_ABSTRACT in t) and "ABSTRACT" not in t.upper())
    en_abs_ctx = _pick_context("abstract", lambda t: "ABSTRACT" in t.upper())
    conclusion_ctx = _pick_context("body", lambda t: "结束语" in t or "结语" in t or "结论" in t)
    thanks_ctx = _pick_context("thanks", lambda t: KW_THANKS in t or KW_ACK in t or "ACKNOWLEDG" in t.upper())
    ref_ctx = _pick_context("references", lambda t: KW_REFERENCES in t or "REFERENCES" in t.upper())
    appendix_ctx = _pick_context("appendix", lambda t: t.startswith(KW_APPENDIX) or t.upper().startswith("APPENDIX"))

    _check("ABSTRACT_ZH_TITLE", "abstract zh", zh_abs_ctx, "abstract_zh")
    _check("ABSTRACT_EN_TITLE", "abstract en", en_abs_ctx, "abstract_en")
    _check("CONCLUSION_TITLE", "conclusion", conclusion_ctx, "conclusion")
    _check("THANKS_TITLE", "thanks", thanks_ctx, "thanks")
    _check("REFERENCE_TITLE", "reference", ref_ctx, "references")
    _check("APPENDIX_TITLE", "appendix", appendix_ctx, "appendix")

    return problems


def _detect_body(contexts: List[Dict], standards: Dict, checks: Dict[str, bool], severities: Dict[str, str]) -> List[Dict]:
    problems = []
    body = standards["body"]
    target_indent = float(body["font_size_pt"]) * float(body["first_line_indent_chars"])
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    for ctx in contexts:
        if ctx["region"] != "body" or ctx["paragraph_type"] != "body":
            continue
        if str(ctx.get("issue_status", "open")).lower() in {"manual_review", "suspected"}:
            continue
        ctx_reasons = [str(item).lower() for item in ctx.get("reasons", [])]
        if "suspected_heading" in ctx_reasons:
            continue

        paragraph = ctx["paragraph"]
        text = ctx["text"]
        pos = f"paragraph {ctx['index']} ({text[:20]})"

        if checks.get("BODY_FONT", True):
            font_name = _get_paragraph_font_name(paragraph)
            if font_name in (None, "unknown"):
                problems.append(
                    _manual_review_problem(
                        "body_font_unknown",
                        pos,
                        "check body font manually",
                        ctx["confidence"],
                        ctx["reasons"] + ["effective_font_name_unknown"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )
            elif font_name != body["font_name"]:
                problems.append(
                    _problem(
                        "body_font",
                        "body font mismatch",
                        pos,
                        body["font_name"],
                        font_name,
                        f"use {body['font_name']}",
                        severity=_severity("BODY_FONT", severities, "ERROR"),
                        auto_fix=True,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + ["body font check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

        if checks.get("BODY_FONT_SIZE", True):
            font_size = _get_paragraph_font_size_pt(paragraph)
            if font_size is None:
                problems.append(
                    _manual_review_problem(
                        "body_font_size_unknown",
                        pos,
                        "check body font size manually",
                        ctx["confidence"],
                        ctx["reasons"] + ["effective_font_size_unknown"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )
            elif abs(font_size - float(body["font_size_pt"])) > 0.1:
                problems.append(
                    _problem(
                        "body_font_size",
                        "body font size mismatch",
                        pos,
                        f"{body['font_size_pt']}pt",
                        f"{font_size:.1f}pt",
                        f"use {body['font_size_pt']}pt",
                        severity=_severity("BODY_FONT_SIZE", severities, "ERROR"),
                        auto_fix=True,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + ["body font size check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

        if checks.get("BODY_LINE_SPACING", True):
            line_spacing = get_effective_line_spacing(paragraph)
            if line_spacing == "unknown":
                problems.append(
                    _manual_review_problem(
                        "body_line_spacing_unknown",
                        pos,
                        "check body line spacing manually",
                        ctx["confidence"],
                        ctx["reasons"] + ["effective_line_spacing_unknown"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )
            elif abs(float(line_spacing) - float(body["line_spacing"])) > 0.05:
                problems.append(
                    _problem(
                        "line_spacing",
                        "body line spacing mismatch",
                        pos,
                        f"{body['line_spacing']}",
                        f"{line_spacing}",
                        f"use {body['line_spacing']}",
                        severity=_severity("BODY_LINE_SPACING", severities, "WARNING"),
                        auto_fix=True,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + ["body line spacing check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

        if checks.get("BODY_FIRST_INDENT", True):
            indent = get_effective_first_line_indent(paragraph)
            if indent == "unknown":
                problems.append(
                    _manual_review_problem(
                        "body_first_indent_unknown",
                        pos,
                        "check body first-line indent manually",
                        ctx["confidence"],
                        ctx["reasons"] + ["effective_first_line_indent_unknown"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )
            elif abs(float(indent) - target_indent) > 1:
                problems.append(
                    _problem(
                        "first_line_indent",
                        "body first line indent mismatch",
                        pos,
                        f"{body['first_line_indent_chars']} chars",
                        f"{float(indent):.1f}pt",
                        f"use {body['first_line_indent_chars']} chars",
                        severity=_severity("BODY_FIRST_INDENT", severities, "WARNING"),
                        auto_fix=True,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + ["body first line indent check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

        if checks.get("BODY_ALIGNMENT", True):
            expected_align_name = str(body.get("align", "justify")).strip().lower()
            expected_align = align_map.get(expected_align_name)
            effective_align = get_effective_alignment(paragraph)
            if effective_align == "unknown":
                problems.append(
                    _manual_review_problem(
                        "body_alignment_unknown",
                        pos,
                        "check body alignment manually",
                        ctx["confidence"],
                        ctx["reasons"] + ["effective_alignment_unknown"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )
            elif expected_align is not None and effective_align != expected_align:
                problems.append(
                    _problem(
                        "body_alignment",
                        "body alignment mismatch",
                        pos,
                        expected_align_name,
                        str(effective_align),
                        f"use {expected_align_name}",
                        severity=_severity("BODY_ALIGNMENT", severities, "WARNING"),
                        auto_fix=True,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + ["body alignment check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

        if checks.get("BODY_PARAGRAPH_SPACING", True):
            before = _pt(paragraph.paragraph_format.space_before) or 0.0
            after = _pt(paragraph.paragraph_format.space_after) or 0.0
            std_before = float(body.get("space_before_pt", 0.0))
            std_after = float(body.get("space_after_pt", 0.0))
            if abs(before - std_before) > 0.5 or abs(after - std_after) > 0.5:
                problems.append(
                    _problem(
                        "paragraph_spacing",
                        "body paragraph spacing mismatch",
                        pos,
                        f"before {std_before}pt / after {std_after}pt",
                        f"before {before:.1f}pt / after {after:.1f}pt",
                        "fix paragraph spacing",
                        severity=_severity("BODY_PARAGRAPH_SPACING", severities, "INFO"),
                        auto_fix=True,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + ["body paragraph spacing check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

    return problems


def _extract_figure_table_numbers(contexts: List[Dict]):
    figures = []
    tables = []
    for ctx in contexts:
        text = ctx["text"]
        if ctx["paragraph_type"] == "figure_caption":
            match = re.match(rf"^({KW_FIGURE}|figure)\s*([0-9]+)", text, flags=re.IGNORECASE)
            if match:
                figures.append((ctx["index"], int(match.group(2)), ctx))
        if ctx["paragraph_type"] == "table_caption":
            match = re.match(rf"^({KW_TABLE}|table)\s*([0-9]+)", text, flags=re.IGNORECASE)
            if match:
                tables.append((ctx["index"], int(match.group(2)), ctx))
    return figures, tables


def _detect_figure_table(
    doc: Document,
    contexts: List[Dict],
    standards: Dict,
    checks: Dict[str, bool],
    severities: Dict[str, str],
) -> List[Dict]:
    if not checks.get("FIGURE_TABLE", True):
        return []

    problems = []
    cfg = standards.get("figure_table", {}) or {}
    check_number_continuity = bool(cfg.get("check_number_continuity", True))
    check_caption_position = bool(cfg.get("check_caption_position", True))
    check_image_center = bool(cfg.get("check_image_center", True))
    figures, tables = _extract_figure_table_numbers(contexts)

    if check_number_continuity:
        for label, values in (("figure", figures), ("table", tables)):
            for idx in range(1, len(values)):
                expected = values[idx - 1][1] + 1
                current = values[idx][1]
                if current != expected:
                    ctx = values[idx][2]
                    problems.append(
                        _problem(
                            "figure_table_number",
                            f"{label} index not continuous",
                            f"paragraph {values[idx][0]}",
                            str(expected),
                            str(current),
                            f"fix {label} index continuity",
                            severity=_severity("FIGURE_TABLE", severities, "WARNING"),
                            auto_fix=False,
                            confidence=ctx["confidence"],
                            reasons=ctx["reasons"] + [f"{label} sequence check"],
                            region=ctx["region"],
                            paragraph_type=ctx["paragraph_type"],
                        )
                    )

    if check_caption_position:
        for para_idx, _, ctx in figures:
            prev_has_image = para_idx > 1 and "graphicData" in doc.paragraphs[para_idx - 2]._p.xml
            if not prev_has_image:
                problems.append(
                    _problem(
                        "figure_caption_position",
                        "figure caption may not be below image",
                        f"paragraph {para_idx}",
                        "caption below image",
                        "previous paragraph has no image",
                        "move caption below corresponding image",
                        severity=_severity("FIGURE_TABLE", severities, "WARNING"),
                        auto_fix=False,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + ["figure caption position check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

        for para_idx, _, ctx in tables:
            next_is_table = False
            if para_idx < len(doc.paragraphs):
                next_node = doc.paragraphs[para_idx]._p.getnext()
                if next_node is not None and next_node.tag.endswith("tbl"):
                    next_is_table = True
            if not next_is_table:
                problems.append(
                    _problem(
                        "table_caption_position",
                        "table caption may not be above table",
                        f"paragraph {para_idx}",
                        "caption above table",
                        "next block is not a table",
                        "move caption above corresponding table",
                        severity=_severity("FIGURE_TABLE", severities, "WARNING"),
                        auto_fix=False,
                        confidence=ctx["confidence"],
                        reasons=ctx["reasons"] + ["table caption position check"],
                        region=ctx["region"],
                        paragraph_type=ctx["paragraph_type"],
                    )
                )

    if check_image_center:
        for ctx in contexts:
            paragraph = ctx["paragraph"]
            if "graphicData" in paragraph._p.xml:
                align = paragraph.paragraph_format.alignment
                if align is not None and align != WD_ALIGN_PARAGRAPH.CENTER:
                    problems.append(
                        _problem(
                            "image_alignment",
                            "image not centered",
                            f"paragraph {ctx['index']}",
                            "center",
                            str(align),
                            "set image paragraph alignment to center",
                            severity=_severity("FIGURE_TABLE", severities, "WARNING"),
                            auto_fix=True,
                            confidence=0.9,
                            reasons=["image alignment check"],
                            region=ctx["region"],
                            paragraph_type=ctx["paragraph_type"],
                        )
                    )

    return problems


def _detect_toc(doc: Document, contexts: List[Dict], standards: Dict, checks: Dict[str, bool], severities: Dict[str, str]) -> List[Dict]:
    if not checks.get("TOC_CHECK", True):
        return []
    toc_cfg = standards.get("toc", {}) or {}
    if not bool(toc_cfg.get("require_auto_toc", True)):
        return []
    has_catalog_region = any(ctx["region"] == "catalog" for ctx in contexts if ctx["paragraph_type"] != "empty")
    has_toc_field = "TOC " in doc._element.xml
    if has_catalog_region and has_toc_field:
        return []
    return [
        _problem(
            "toc",
            "toc field or catalog region not detected",
            "document",
            "auto-generated toc",
            "not detected",
            "generate catalog by Word TOC field",
            severity=_severity("TOC_CHECK", severities, "WARNING"),
            auto_fix=False,
            confidence=0.8,
            reasons=["catalog region / TOC field check"],
            region="catalog" if has_catalog_region else "unknown",
            paragraph_type="section_marker" if has_catalog_region else "unknown",
        )
    ]


def _extract_references(contexts: List[Dict]):
    refs = []
    for ctx in contexts:
        if ctx["region"] != "references":
            continue
        if ctx["paragraph_type"] == "section_marker":
            continue
        if not ctx["text"]:
            continue
        refs.append((ctx["index"], ctx["text"], ctx))
    return refs


def _extract_reference_region_entries(contexts: List[Dict]):
    entries = []
    for ctx in contexts:
        if ctx["region"] != "references":
            continue
        if ctx["paragraph_type"] == "section_marker":
            continue
        text = (ctx.get("text") or "").strip()
        if not text:
            continue
        # Guard against region boundary errors: only treat likely bibliography
        # lines as reference entries, so normal body paragraphs are not
        # force-checked as references.
        if ctx.get("paragraph_type") == "reference_item" or _is_reference_item(text):
            entries.append((ctx["index"], text, ctx))
    return entries


def _extract_tail_reference_candidates(contexts: List[Dict], tail_ratio: float = 0.3):
    if not contexts:
        return []
    start = max(1, int(len(contexts) * (1.0 - tail_ratio)))
    candidates = []
    for ctx in contexts[start - 1 :]:
        if not ctx["text"]:
            continue
        if _is_reference_item(ctx["text"]):
            candidates.append((ctx["index"], ctx["text"], ctx))
    return candidates


def _parse_reference_number(text: str):
    value = (text or "").strip()
    m = re.match(r"^\[(\d+)\]\s*(.*)$", value)
    if m:
        return int(m.group(1)), "bracket", m.group(2).strip()
    m = re.match(r"^(\d+)\.\s*(.*)$", value)
    if m:
        return int(m.group(1)), "dot", m.group(2).strip()
    m = re.match(r"^(\d+)\u3001\s*(.*)$", value)
    if m:
        return int(m.group(1)), "cn_comma", m.group(2).strip()
    return None, None, value


def _detect_references(contexts: List[Dict], standards: Dict, checks: Dict[str, bool], severities: Dict[str, str]) -> List[Dict]:
    problems = []
    ref_cfg = standards.get("references", {}) or {}
    check_completeness = bool(ref_cfg.get("check_completeness", True))
    check_type_mark = bool(ref_cfg.get("check_type_mark", True))
    region_markers = [ctx for ctx in contexts if ctx["region"] == "references" and ctx["paragraph_type"] == "section_marker"]
    has_reference_region = len(region_markers) > 0
    region_entries = _extract_reference_region_entries(contexts)

    if not has_reference_region:
        tail_candidates = _extract_tail_reference_candidates(contexts)
        if tail_candidates:
            problems.append(
                _problem(
                    "reference_region_uncertain",
                    "possible references detected near document end, but reference region marker not found",
                    "document",
                    "explicit references region marker",
                    "not found",
                    "manual review: confirm reference section heading",
                    severity="INFO",
                    auto_fix=False,
                    confidence=0.58,
                    issue_status="manual_review",
                    reasons=["no_reference_marker", "numbered_entries_near_end"],
                    region="unknown",
                    paragraph_type="reference_item",
                )
            )
        else:
            problems.append(
                _problem(
                    "reference_region_missing",
                    "reference region not detected",
                    "document",
                    "references region exists",
                    "not detected",
                    "manual review: confirm whether references section is present",
                    severity="INFO",
                    auto_fix=False,
                    confidence=0.5,
                    issue_status="manual_review",
                    reasons=["no_reference_marker", "no_tail_reference_candidates"],
                    region="unknown",
                    paragraph_type="unknown",
                )
            )
        return problems

    if not region_entries:
        problems.append(
            _problem(
                "reference_entries_empty",
                "references region exists but contains no entries",
                "references region",
                "at least one reference entry",
                "0",
                "add reference entries under references heading",
                severity=_severity("REFERENCE_FORMAT", severities, "WARNING"),
                auto_fix=False,
                confidence=0.92,
                reasons=["reference_region_detected", "no_nonempty_entries"],
                region="references",
                paragraph_type="reference_item",
            )
        )
        return problems

    parsed = []
    formats = set()
    for idx, text, ctx in region_entries:
        number, fmt, body = _parse_reference_number(text)
        if fmt is None:
            problems.append(
                _problem(
                    "reference_entry_unrecognized",
                    "entry numbering format not recognized",
                    f"paragraph {idx} ({text[:20]})",
                    "supported: [n] / n. / n\u3001",
                    text[:40],
                    "manual review: verify if this paragraph is a reference entry",
                    severity="INFO",
                    auto_fix=False,
                    confidence=0.6,
                    issue_status="manual_review",
                    reasons=ctx["reasons"] + ["unsupported_reference_number_format"],
                    region="references",
                    paragraph_type=ctx["paragraph_type"],
                )
            )
            continue

        formats.add(fmt)
        parsed.append((idx, number, fmt, body, ctx))
        if checks.get("REFERENCE_FORMAT", True) and check_completeness and not body:
            problems.append(
                _problem(
                    "reference_entry_empty",
                    "reference entry content is empty",
                    f"paragraph {idx} ({text[:20]})",
                    "entry has content after number",
                    "empty",
                    "fill this reference entry content",
                    severity=_severity("REFERENCE_FORMAT", severities, "WARNING"),
                    auto_fix=False,
                    confidence=0.96,
                    reasons=ctx["reasons"] + ["empty_reference_entry"],
                    region="references",
                    paragraph_type=ctx["paragraph_type"],
                )
            )
        elif checks.get("REFERENCE_FORMAT", True) and check_type_mark and not re.search(r"\[[A-Za-z/]+\]", body or ""):
            problems.append(
                _problem(
                    "reference_type_mark",
                    "reference type mark missing or invalid",
                    f"paragraph {idx} ({text[:20]})",
                    "contains [J]/[M]/[D]/[C]/[R]/[S]/[P]/[EB/OL] etc.",
                    body[:40] if isinstance(body, str) else "",
                    "add standard reference type mark",
                    severity=_severity("REFERENCE_FORMAT", severities, "WARNING"),
                    auto_fix=False,
                    confidence=0.88,
                    reasons=ctx["reasons"] + ["reference_type_mark_check"],
                    region="references",
                    paragraph_type=ctx["paragraph_type"],
                )
            )

    if not parsed:
        problems.append(
            _problem(
                "reference_entries_uncertain",
                "references region found, but no entries matched supported numbering format",
                "references region",
                "supported numbering entries",
                "none",
                "manual review: use [n], n. or n\u3001 format",
                severity="INFO",
                auto_fix=False,
                confidence=0.66,
                issue_status="manual_review",
                reasons=["reference_region_detected", "no_supported_numbering_entries"],
                region="references",
                paragraph_type="reference_item",
            )
        )
        return problems

    if len(formats) > 1:
        problems.append(
            _problem(
                "reference_numbering_mixed_format",
                "multiple numbering formats are mixed in references",
                "references region",
                "single numbering format",
                ", ".join(sorted(formats)),
                "use one consistent numbering format",
                severity=_severity("REFERENCE_SEQ", severities, "WARNING"),
                auto_fix=False,
                confidence=0.9,
                reasons=["mixed_reference_numbering_format"],
                region="references",
                paragraph_type="reference_item",
            )
        )

    if checks.get("REFERENCE_SEQ", True):
        numbers = [item[1] for item in parsed]
        index_by_number: Dict[int, List[int]] = {}
        for para_idx, num, _, _, _ in parsed:
            index_by_number.setdefault(num, []).append(para_idx)

        # duplicate
        for num, para_list in index_by_number.items():
            if len(para_list) > 1:
                problems.append(
                    _problem(
                        "reference_number_duplicate",
                        "duplicate reference number detected",
                        f"paragraph {para_list[0]}",
                        "unique numbering",
                        f"{num} appears {len(para_list)} times",
                        "remove or renumber duplicate entries",
                        severity=_severity("REFERENCE_SEQ", severities, "ERROR"),
                        auto_fix=False,
                        confidence=0.97,
                        reasons=["duplicate_reference_number"],
                        region="references",
                        paragraph_type="reference_item",
                    )
                )

        # out of order
        for i in range(1, len(parsed)):
            prev_num = parsed[i - 1][1]
            cur_num = parsed[i][1]
            if cur_num < prev_num:
                problems.append(
                    _problem(
                        "reference_number_out_of_order",
                        "reference numbers are out of order",
                        f"paragraph {parsed[i][0]}",
                        f">= {prev_num}",
                        str(cur_num),
                        "sort or renumber reference entries",
                        severity=_severity("REFERENCE_SEQ", severities, "ERROR"),
                        auto_fix=False,
                        confidence=0.96,
                        reasons=["reference_number_order_check"],
                        region="references",
                        paragraph_type="reference_item",
                    )
                )

        # continuity
        sorted_unique = sorted(set(numbers))
        if sorted_unique:
            expected = list(range(sorted_unique[0], sorted_unique[-1] + 1))
            missing = [num for num in expected if num not in set(sorted_unique)]
            if missing:
                problems.append(
                    _problem(
                        "reference_number_not_continuous",
                        "reference numbering is not continuous",
                        "references region",
                        f"continuous numbers: {expected[0]}..{expected[-1]}",
                        f"missing: {missing}",
                        "renumber references continuously",
                        severity=_severity("REFERENCE_SEQ", severities, "ERROR"),
                        auto_fix=False,
                        confidence=0.95,
                        reasons=["reference_number_continuity_check"],
                        region="references",
                        paragraph_type="reference_item",
                    )
                )

    return problems


def _build_structure_summary(contexts: List[Dict]) -> Dict:
    region_stats = {region: 0 for region in REGIONS}
    type_stats = {para_type: 0 for para_type in PARAGRAPH_TYPES}
    for ctx in contexts:
        region_stats[ctx["region"]] = region_stats.get(ctx["region"], 0) + 1
        type_stats[ctx["paragraph_type"]] = type_stats.get(ctx["paragraph_type"], 0) + 1
    return {"region_stats": region_stats, "paragraph_type_stats": type_stats}


def _normalize_issue_status_for_scoring(issue: Dict) -> str:
    status = str(issue.get("issue_status", "") or "").strip().lower()
    confidence = float(issue.get("confidence", 0.0) or 0.0)
    reasons = [str(item).lower() for item in issue.get("reasons", [])]

    if status in {"confirmed", "suspected", "manual_review"}:
        return status
    if status == "open":
        if any("suspected" in item for item in reasons):
            return "suspected"
        return "confirmed" if confidence >= 0.68 else "manual_review"
    if not status:
        return "confirmed" if confidence >= 0.68 else "manual_review"
    return "manual_review"


def detect_word_format(
    docx_path: str,
    standards: Optional[Dict] = None,
    checks: Optional[Dict[str, bool]] = None,
    severities: Optional[Dict[str, str]] = None,
) -> Dict:
    doc = Document(docx_path)
    final_standards = _merge(DEFAULT_STANDARDS, standards or {})
    final_checks = dict(DEFAULT_CHECKS)
    if checks:
        final_checks.update(checks)
    final_severities = severities or {}

    # Pipeline:
    # 1) read docx
    # 2) detect document regions
    # 3) detect paragraph types
    # 4) apply rule checks by region/type
    # 5) build issues + score
    contexts = _analyze_document_structure(doc, final_standards)

    issues: List[Dict] = []
    issues.extend(_detect_page(doc, final_standards, final_checks, final_severities, contexts=contexts))
    issues.extend(_detect_headings(contexts, final_standards, final_checks, final_severities))
    issues.extend(_detect_section_titles(contexts, final_standards, final_checks, final_severities))
    issues.extend(_detect_body(contexts, final_standards, final_checks, final_severities))
    issues.extend(_detect_figure_table(doc, contexts, final_standards, final_checks, final_severities))
    issues.extend(_detect_toc(doc, contexts, final_standards, final_checks, final_severities))
    issues.extend(_detect_references(contexts, final_standards, final_checks, final_severities))

    error_count = sum(1 for item in issues if item["severity"] == "ERROR")
    warning_count = sum(1 for item in issues if item["severity"] == "WARNING")
    info_count = sum(1 for item in issues if item["severity"] == "INFO")

    confirmed_error_count = 0
    confirmed_warning_count = 0
    confirmed_info_count = 0
    suspected_error_count = 0
    suspected_warning_count = 0
    suspected_info_count = 0
    manual_review_count = 0

    for item in issues:
        status = _normalize_issue_status_for_scoring(item)
        sev = item.get("severity")
        if status == "manual_review":
            manual_review_count += 1
            continue
        if status == "suspected":
            if sev == "ERROR":
                suspected_error_count += 1
            elif sev == "WARNING":
                suspected_warning_count += 1
            elif sev == "INFO":
                suspected_info_count += 1
            continue
        if sev == "ERROR":
            confirmed_error_count += 1
        elif sev == "WARNING":
            confirmed_warning_count += 1
        elif sev == "INFO":
            confirmed_info_count += 1

    penalty = (
        min(40.0, confirmed_error_count * 2.0)
        + min(20.0, confirmed_warning_count * 0.5)
        + min(10.0, confirmed_info_count * 0.1)
        + min(15.0, suspected_error_count * 0.5)
        + min(10.0, suspected_warning_count * 0.2)
        + min(5.0, suspected_info_count * 0.05)
    )
    score = round(max(0.0, 100.0 - penalty), 2)

    return {
        "score": score,
        "pass_flag": 1 if score >= 80.0 else 0,
        "error_count": error_count,
        "warning_count": warning_count,
        "info_count": info_count,
        "confirmed_error_count": confirmed_error_count,
        "confirmed_warning_count": confirmed_warning_count,
        "confirmed_info_count": confirmed_info_count,
        "suspected_error_count": suspected_error_count,
        "suspected_warning_count": suspected_warning_count,
        "suspected_info_count": suspected_info_count,
        "manual_review_count": manual_review_count,
        "hit_rule_count": len(issues),
        "issues": issues,
        "document_structure": _build_structure_summary(contexts),
    }


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Word format checker")
    parser.add_argument("docx_path", help="docx file path")
    parser.add_argument("--standards", default=None, help="standards json")
    parser.add_argument("--checks", default=None, help="checks json")
    args = parser.parse_args()

    standards_arg = json.loads(args.standards) if args.standards else None
    checks_arg = json.loads(args.checks) if args.checks else None
    result = detect_word_format(args.docx_path, standards=standards_arg, checks=checks_arg)
    print(json.dumps(result, ensure_ascii=False, indent=2))
