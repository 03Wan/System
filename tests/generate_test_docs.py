import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SAMPLES_DIR = ROOT / "samples"
EXPECTED_DIR = ROOT / "expected"

FONT_SIMSUN = "\u5b8b\u4f53"
FONT_HEITI = "\u9ed1\u4f53"


def _set_run_font(run, name: str, size: float, bold: bool = False):
    run.font.name = name
    if run._element is not None and run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def _add_heading(doc: Document, text: str, level: int, bad: bool = False):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(0)
    if level == 1:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 14 if bad else 16
    elif level == 2:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 12 if bad else 14
    else:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 11 if bad else 12
    run = p.runs[0]
    _set_run_font(run, FONT_SIMSUN if bad else FONT_HEITI, size, bold=not bad)
    return p


def _add_body_paragraph(
    doc: Document,
    text: str,
    bad_font: bool = False,
    bad_size: bool = False,
    bad_spacing: bool = False,
    bad_indent: bool = False,
):
    p = doc.add_paragraph(text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0 if bad_spacing else 1.5
    p.paragraph_format.first_line_indent = Pt(0) if bad_indent else Pt(24)
    run = p.runs[0]
    font_name = "Arial" if bad_font else FONT_SIMSUN
    font_size = 14 if bad_size else 12
    _set_run_font(run, font_name, font_size, bold=False)
    return p


def _add_references(doc: Document, items):
    _add_heading(doc, "\u53c2\u8003\u6587\u732e", level=1, bad=False)
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.5
        if p.runs:
            _set_run_font(p.runs[0], FONT_SIMSUN, 10.5, bold=False)


def _build_correct_doc(path: Path):
    doc = Document()
    _add_heading(doc, "\u7b2c\u4e00\u7ae0 \u7eea\u8bba", level=1, bad=False)
    _add_body_paragraph(doc, "\u8fd9\u662f\u4e00\u6bb5\u6b63\u5e38\u6b63\u6587\uff0c\u7528\u4e8e\u68c0\u9a8c\u6837\u672c\u683c\u5f0f\u6b63\u786e\u3002")
    _add_heading(doc, "1.1 \u7814\u7a76\u80cc\u666f", level=2, bad=False)
    _add_body_paragraph(doc, "\u8fd9\u662f\u53e6\u4e00\u6bb5\u6b63\u5e38\u6b63\u6587\uff0c\u5e94\u8be5\u4e0d\u4ea7\u751f\u6b63\u6587\u683c\u5f0f\u9519\u8bef\u3002")
    _add_references(
        doc,
        [
            "[1] Zhang San. Sample Paper A. Journal A, 2020.",
            "[2] Li Si. Sample Paper B. Journal B, 2021.",
        ],
    )
    doc.save(path)


def _write_expected(path: Path, expected_problem_types):
    payload = {
        "sample": path.name.replace(".json", ".docx"),
        "expected_problem_types": expected_problem_types,
        "min_expected_hits": len(expected_problem_types),
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main():
    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    EXPECTED_DIR.mkdir(parents=True, exist_ok=True)

    # 1. correct.docx
    correct = SAMPLES_DIR / "correct.docx"
    _build_correct_doc(correct)
    _write_expected(EXPECTED_DIR / "correct.json", [])

    # 2. bad_body_font.docx
    bad_body_font = SAMPLES_DIR / "bad_body_font.docx"
    doc = Document()
    _add_heading(doc, "\u7b2c\u4e00\u7ae0 \u7eea\u8bba", level=1)
    _add_body_paragraph(doc, "\u6b63\u6587\u5b57\u4f53\u9519\u8bef\u6837\u672c\u3002", bad_font=True)
    _add_references(doc, ["[1] A. Ref One. Journal, 2020."])
    doc.save(bad_body_font)
    _write_expected(EXPECTED_DIR / "bad_body_font.json", ["body_font"])

    # 3. bad_body_size.docx
    bad_body_size = SAMPLES_DIR / "bad_body_size.docx"
    doc = Document()
    _add_heading(doc, "\u7b2c\u4e00\u7ae0 \u7eea\u8bba", level=1)
    _add_body_paragraph(doc, "\u6b63\u6587\u5b57\u53f7\u9519\u8bef\u6837\u672c\u3002", bad_size=True)
    _add_references(doc, ["[1] A. Ref One. Journal, 2020."])
    doc.save(bad_body_size)
    _write_expected(EXPECTED_DIR / "bad_body_size.json", ["body_font_size"])

    # 4. bad_line_spacing.docx
    bad_line_spacing = SAMPLES_DIR / "bad_line_spacing.docx"
    doc = Document()
    _add_heading(doc, "\u7b2c\u4e00\u7ae0 \u7eea\u8bba", level=1)
    _add_body_paragraph(doc, "\u6b63\u6587\u884c\u8ddd\u9519\u8bef\u6837\u672c\u3002", bad_spacing=True)
    _add_references(doc, ["[1] A. Ref One. Journal, 2020."])
    doc.save(bad_line_spacing)
    _write_expected(EXPECTED_DIR / "bad_line_spacing.json", ["line_spacing"])

    # 5. bad_indent.docx
    bad_indent = SAMPLES_DIR / "bad_indent.docx"
    doc = Document()
    _add_heading(doc, "\u7b2c\u4e00\u7ae0 \u7eea\u8bba", level=1)
    _add_body_paragraph(doc, "\u6b63\u6587\u9996\u884c\u7f29\u8fdb\u9519\u8bef\u6837\u672c\u3002", bad_indent=True)
    _add_references(doc, ["[1] A. Ref One. Journal, 2020."])
    doc.save(bad_indent)
    _write_expected(EXPECTED_DIR / "bad_indent.json", ["first_line_indent"])

    # 6. bad_heading.docx
    bad_heading = SAMPLES_DIR / "bad_heading.docx"
    doc = Document()
    _add_heading(doc, "\u7b2c\u4e00\u7ae0 \u7eea\u8bba", level=1, bad=True)
    _add_body_paragraph(doc, "\u7528\u4e8e\u914d\u5408\u9519\u8bef\u6807\u9898\u7684\u6b63\u6587\u6bb5\u843d\u3002")
    _add_references(doc, ["[1] A. Ref One. Journal, 2020."])
    doc.save(bad_heading)
    _write_expected(EXPECTED_DIR / "bad_heading.json", ["heading_format"])

    # 7. bad_reference_number.docx
    bad_ref = SAMPLES_DIR / "bad_reference_number.docx"
    doc = Document()
    _add_heading(doc, "\u7b2c\u4e00\u7ae0 \u7eea\u8bba", level=1)
    _add_body_paragraph(doc, "\u7528\u4e8e\u53c2\u8003\u6587\u732e\u7f16\u53f7\u68c0\u6d4b\u3002")
    _add_references(
        doc,
        [
            "[1] A. Ref One. Journal, 2020.",
            "[3] B. Ref Two. Journal, 2021.",
        ],
    )
    doc.save(bad_ref)
    _write_expected(EXPECTED_DIR / "bad_reference_number.json", ["reference_number_not_continuous"])

    # 8. bad_mixed.docx
    bad_mixed = SAMPLES_DIR / "bad_mixed.docx"
    doc = Document()
    _add_heading(doc, "\u7b2c\u4e00\u7ae0 \u7eea\u8bba", level=1, bad=True)
    _add_body_paragraph(doc, "\u6df7\u5408\u9519\u8bef\u6837\u672c\u7684\u6b63\u6587\uff0c\u5b57\u4f53\u9519\u8bef\u3002", bad_font=True)
    _add_body_paragraph(doc, "\u6df7\u5408\u9519\u8bef\u6837\u672c\u7684\u6b63\u6587\uff0c\u9996\u884c\u7f29\u8fdb\u9519\u8bef\u3002", bad_indent=True)
    _add_references(
        doc,
        [
            "[1] A. Ref One. Journal, 2020.",
            "[3] B. Ref Two. Journal, 2021.",
        ],
    )
    doc.save(bad_mixed)
    _write_expected(
        EXPECTED_DIR / "bad_mixed.json",
        ["heading_format", "body_font", "first_line_indent", "reference_number_not_continuous"],
    )

    # 9. bad_heading_no_space.docx
    # Heading number has no following space, which is common in real docs (e.g. "1.1研究背景").
    # It should still be recognized as a heading and checked against heading style rules.
    bad_heading_no_space = SAMPLES_DIR / "bad_heading_no_space.docx"
    doc = Document()
    _add_heading(doc, "第一章 绪论", level=1, bad=False)
    p = doc.add_paragraph("1.1研究背景")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    if p.runs:
        _set_run_font(p.runs[0], FONT_SIMSUN, 12, bold=False)  # intentionally wrong for level-2 heading
    _add_body_paragraph(doc, "用于测试无空格编号标题识别与格式检测。")
    _add_references(doc, ["[1] A. Ref One. Journal, 2020."])
    doc.save(bad_heading_no_space)
    _write_expected(EXPECTED_DIR / "bad_heading_no_space.json", ["heading_format"])

    # 10. list_item_body.docx
    # Numbered body content should not be treated as a heading.
    list_item_body = SAMPLES_DIR / "list_item_body.docx"
    doc = Document()
    _add_heading(doc, "第一章 绪论", level=1, bad=False)
    _add_body_paragraph(doc, "1. 这是一个编号句子，用于描述步骤，不应被判定为标题。")
    _add_body_paragraph(doc, "该段继续说明内容，保持正文格式。")
    _add_references(doc, ["[1] A. Ref One. Journal, 2020."])
    doc.save(list_item_body)
    _write_expected(EXPECTED_DIR / "list_item_body.json", [])

    print(f"Generated samples in: {SAMPLES_DIR}")
    print(f"Generated expected JSON in: {EXPECTED_DIR}")


if __name__ == "__main__":
    main()
